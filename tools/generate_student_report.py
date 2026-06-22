from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"D:\LTDD\android_todoapp-main")
OUT = ROOT / "Bao_cao_Student_Planner_LTDD_Hoan_Chinh.docx"
ASSET_DIR = ROOT / "report_assets"
ASSET_DIR.mkdir(exist_ok=True)

FONT_TIMES = "Times New Roman"
DARK = RGBColor(0, 0, 0)
GRAY = RGBColor(90, 90, 90)
LIGHT_BLUE = "D9EAF7"


def set_run_font(run, name=FONT_TIMES, size=13, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def para_format(p, before=0, after=0, line=1.5, align=None, first_line=True):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if first_line:
        pf.first_line_indent = Inches(0.3)
    if align is not None:
        p.alignment = align


def add_p(doc, text="", bold=False, italic=False, size=13, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
          before=0, after=0, line=1.5, first_line=True, color=DARK):
    p = doc.add_paragraph()
    para_format(p, before, after, line, align, first_line)
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, italic=italic, color=color)
    return p


def add_center(doc, text, size=13, bold=False, after=0, before=0):
    return add_p(doc, text, bold=bold, size=size, align=WD_ALIGN_PARAGRAPH.CENTER,
                 before=before, after=after, line=1.15, first_line=False)


def add_h1(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text.upper())
    set_run_font(r, size=15, bold=True)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.first_line_indent = Inches(0)
    r = p.add_run(text)
    set_run_font(r, size=13, bold=True)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph(style="Heading 3")
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.first_line_indent = Inches(0)
    r = p.add_run(text)
    set_run_font(r, size=13, bold=True, italic=True)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r, size=13)
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, size=11.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(str(text))
    set_run_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell_margins(cell)


def set_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)


def add_table(doc, headers, rows, widths=None, font_size=11.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if widths:
        set_widths(table, widths)
    for i, h in enumerate(headers):
        shade_cell(table.rows[0].cells[i], LIGHT_BLUE)
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=font_size,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=font_size)
    add_p(doc, "", after=4, line=1, first_line=False)
    return table


def add_caption(doc, text):
    return add_p(doc, text, italic=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER,
                 before=2, after=8, line=1.15, first_line=False, color=GRAY)


def page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, sep, text, end])


def configure(doc):
    for section in doc.sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(0.79)
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.5)
    normal = doc.styles["Normal"]
    normal.font.name = FONT_TIMES
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_TIMES)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_TIMES)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_TIMES)
    normal.font.size = Pt(13)
    normal.paragraph_format.line_spacing = 1.5
    for style_name in ["Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"]:
        st = doc.styles[style_name]
        st.font.name = FONT_TIMES
        st._element.rPr.rFonts.set(qn("w:ascii"), FONT_TIMES)
        st._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_TIMES)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_TIMES)
        st.paragraph_format.page_break_before = False


def set_footer(doc):
    for section in doc.sections:
        p = section.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("Student Planner - Lập trình thiết bị di động | Trang ")
        set_run_font(r, size=10, color=GRAY)
        page_number(p)


def get_font(size=28, bold=False):
    path = Path(r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf")
    if path.exists():
        return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def wrap(draw, text, font, max_width):
    words = text.split()
    lines, current = [], ""
    for word in words:
        test = word if not current else current + " " + word
        if draw.textbbox((0, 0), test, font=font)[2] <= max_width:
            current = test
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def box(draw, xy, text, fill, outline, font):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
    lines = wrap(draw, text, font, x2 - x1 - 28)
    y = y1 + (y2 - y1 - len(lines) * (font.size + 4)) / 2
    for line in lines:
        w = draw.textbbox((0, 0), line, font=font)[2]
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=font, fill=(20, 20, 20))
        y += font.size + 4


def arrow(draw, start, end, color=(31, 78, 121)):
    draw.line([start, end], fill=color, width=4)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) > abs(ey - sy):
        sign = 1 if ex > sx else -1
        pts = [(ex, ey), (ex - sign * 14, ey - 9), (ex - sign * 14, ey + 9)]
    else:
        sign = 1 if ey > sy else -1
        pts = [(ex, ey), (ex - 9, ey - sign * 14), (ex + 9, ey - sign * 14)]
    draw.polygon(pts, fill=color)


def diagrams():
    title = get_font(34, True)
    font = get_font(24)
    small = get_font(21)

    img = Image.new("RGB", (1500, 900), "white")
    d = ImageDraw.Draw(img)
    d.text((60, 35), "Kiến trúc tổng quan ứng dụng Student Planner", font=title, fill=(31, 78, 121))
    data = [
        ((80, 150, 360, 260), "Login/Register\nSessionManager"),
        ((610, 135, 900, 275), "MainActivity\nBottomNavigation"),
        ((80, 390, 360, 520), "HomeFragment\nDashboard + gợi ý"),
        ((430, 390, 710, 520), "TaskFragment\nTìm kiếm + lọc"),
        ((780, 390, 1060, 520), "CalendarFragment\nNgày/Tuần/Tháng"),
        ((1130, 390, 1410, 520), "ProfileFragment\nTài khoản"),
        ((300, 675, 650, 815), "DatabaseHelper\nSQLite offline"),
        ((850, 675, 1200, 815), "ReminderScheduler\nAlarm + Notification"),
    ]
    for i, (xy, text) in enumerate(data):
        fill = (230, 244, 255) if i != 7 else (222, 248, 245)
        box(d, xy, text, fill, (31, 118, 210), font if i < 2 else small)
    arrow(d, (360, 205), (610, 205))
    for x in [220, 570, 920, 1270]:
        arrow(d, (755, 275), (x, 390))
        arrow(d, (x, 520), (475, 675))
    arrow(d, (920, 520), (1025, 675))
    img.save(ASSET_DIR / "architecture.png")

    img = Image.new("RGB", (1500, 1000), "white")
    d = ImageDraw.Draw(img)
    d.text((60, 35), "Mô hình dữ liệu SQLite", font=title, fill=(31, 78, 121))
    header = get_font(25, True)
    tf = get_font(21)

    def table_img(x, y, w, h, name, fields, fill):
        d.rounded_rectangle((x, y, x + w, y + h), radius=16, fill="white", outline=(120, 120, 120), width=3)
        d.rounded_rectangle((x, y, x + w, y + 55), radius=16, fill=fill, outline=(120, 120, 120), width=3)
        d.text((x + 20, y + 14), name, font=header, fill=(0, 0, 0))
        yy = y + 75
        for f in fields:
            d.text((x + 20, yy), f, font=tf, fill=(30, 30, 30))
            yy += 34

    table_img(80, 155, 360, 300, "users", ["PK user_id", "username", "email", "password", "created_at"], (217, 234, 247))
    table_img(560, 155, 400, 360, "plan_categories", ["PK category_id", "category_name", "category_code", "note", "color", "FK user_id"], (222, 248, 245))
    table_img(1040, 155, 400, 690, "tasks", ["PK task_id", "title, description", "date, time, end_time", "status, priority", "plan_type", "duration_minutes", "reminder_*", "location, room, subject", "repeat_rule, repeat_until", "wage, submitted", "FK category_id", "FK user_id"], (246, 239, 255))
    arrow(d, (440, 305), (560, 305))
    arrow(d, (960, 335), (1040, 335))
    arrow(d, (440, 390), (1040, 740))
    d.text((458, 268), "1 - n", font=tf, fill=(31, 78, 121))
    d.text((978, 298), "1 - n", font=tf, fill=(31, 78, 121))
    d.text((670, 565), "users.user_id -> tasks.user_id", font=small, fill=(31, 78, 121))
    img.save(ASSET_DIR / "erd.png")

    img = Image.new("RGB", (1500, 880), "white")
    d = ImageDraw.Draw(img)
    d.text((60, 35), "Luồng xử lý thêm kế hoạch và nhắc lịch", font=title, fill=(31, 78, 121))
    steps = [
        ((80, 170, 350, 285), "Nhập nội dung\nloại kế hoạch"),
        ((430, 170, 700, 285), "Chọn ngày giờ\nlặp lại/nhắc trước"),
        ((780, 170, 1050, 285), "Kiểm tra dữ liệu\nvà trùng giờ"),
        ((1130, 170, 1400, 285), "Lưu SQLite\nuser_id hiện tại"),
        ((300, 510, 570, 625), "Tạo AlarmManager\nnếu bật nhắc"),
        ((650, 510, 920, 625), "Đến thời điểm\nReminderReceiver"),
        ((1000, 510, 1270, 625), "Hiển thị thông báo\nmở về app"),
    ]
    for idx, (xy, text) in enumerate(steps):
        box(d, xy, text, (230, 244, 255) if idx < 4 else (222, 248, 245), (31, 118, 210), font)
    arrow(d, (350, 227), (430, 227))
    arrow(d, (700, 227), (780, 227))
    arrow(d, (1050, 227), (1130, 227))
    arrow(d, (1265, 285), (435, 510))
    arrow(d, (570, 568), (650, 568))
    arrow(d, (920, 568), (1000, 568))
    d.text((82, 720), "Nếu lịch lặp có ngày bị trùng giờ, ứng dụng bỏ qua ngày trùng và thông báo số kế hoạch đã tạo.", font=small, fill=(70, 70, 70))
    img.save(ASSET_DIR / "workflow.png")


def remove_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                tag = OxmlElement(f"w:{edge}")
                tag.set(qn("w:val"), "nil")
                borders.append(tag)
            tc_pr.append(borders)


def toc(doc):
    add_h1(doc, "MỤC LỤC")
    entries = [
        ("LỜI CẢM ƠN", "2"), ("NHẬN XÉT CỦA GIÁO VIÊN HƯỚNG DẪN", "3"),
        ("DANH MỤC TỪ VIẾT TẮT", "5"), ("DANH MỤC HÌNH ẢNH", "6"),
        ("DANH MỤC BẢNG BIỂU", "6"), ("TÓM TẮT ĐỀ TÀI", "7"),
        ("Chương 1: GIỚI THIỆU ĐỀ TÀI", "8"), ("Chương 2: CƠ SỞ LÝ THUYẾT", "11"),
        ("Chương 3: PHÂN TÍCH, THIẾT KẾ VÀ CÀI ĐẶT", "14"),
        ("Chương 4: KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", "25"),
        ("TÀI LIỆU THAM KHẢO", "27"),
    ]
    for title, page in entries:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.0), WD_ALIGN_PARAGRAPH.RIGHT)
        set_run_font(p.add_run(title), size=12)
        p.add_run("\t")
        set_run_font(p.add_run(page), size=12)


def main():
    diagrams()
    doc = Document()
    configure(doc)

    add_center(doc, "TRƯỜNG ĐẠI HỌC TRÀ VINH", 14, True)
    add_center(doc, "TRƯỜNG KỸ THUẬT VÀ CÔNG NGHỆ", 14, True)
    add_center(doc, "KHOA CÔNG NGHỆ THÔNG TIN", 14, True, after=10)
    add_center(doc, "---------------------", 13, after=55)
    add_center(doc, "BÁO CÁO BÀI TẬP LỚN", 18, True, after=10)
    add_center(doc, "HỌC PHẦN: LẬP TRÌNH THIẾT BỊ DI ĐỘNG", 15, True, after=38)
    add_center(doc, "ĐỀ TÀI:", 15, True, after=8)
    add_center(doc, "XÂY DỰNG ỨNG DỤNG ANDROID", 18, True)
    add_center(doc, "QUẢN LÝ KẾ HOẠCH CHO SINH VIÊN", 18, True, after=58)
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_widths(t, [3.1, 3.1])
    set_cell_text(t.cell(0, 0), "Giáo viên hướng dẫn\nThS. Phạm Thị Trúc Mai", size=13)
    set_cell_text(t.cell(0, 1), "Sinh viên thực hiện:\nHọ và tên SV 1 (MSSV)\nHọ và tên SV 2 (MSSV)\nHọ và tên SV 3 (MSSV)\nMã lớp: ................", size=13)
    remove_table_borders(t)
    add_center(doc, "Vĩnh Long, tháng 06 năm 2026", 13, True, before=70)
    doc.add_page_break()

    add_h1(doc, "LỜI CẢM ƠN")
    add_p(doc, "Nhóm sinh viên xin chân thành cảm ơn quý Thầy Cô Khoa Công nghệ thông tin, Trường Kỹ thuật và Công nghệ, Trường Đại học Trà Vinh đã tạo điều kiện học tập và rèn luyện trong học phần Lập trình thiết bị di động. Đặc biệt, nhóm xin cảm ơn giảng viên hướng dẫn đã định hướng, góp ý và hỗ trợ nhóm trong quá trình phân tích yêu cầu, xây dựng ứng dụng và hoàn thiện báo cáo.")
    add_p(doc, "Trong quá trình thực hiện đề tài, nhóm đã vận dụng kiến thức về lập trình Android, thiết kế giao diện XML, xử lý dữ liệu SQLite, quản lý phiên đăng nhập, thông báo hệ thống và kiểm thử chức năng. Do thời gian và kinh nghiệm còn hạn chế, báo cáo khó tránh khỏi thiếu sót. Nhóm rất mong nhận được ý kiến đóng góp của Thầy Cô để ứng dụng được hoàn thiện hơn trong các phiên bản tiếp theo.")
    add_p(doc, "Nhóm sinh viên thực hiện", align=WD_ALIGN_PARAGRAPH.RIGHT, first_line=False, before=20)
    add_p(doc, "Sinh viên ký và ghi rõ họ tên", align=WD_ALIGN_PARAGRAPH.RIGHT, first_line=False)
    doc.add_page_break()

    add_h1(doc, "NHẬN XÉT CỦA GIÁO VIÊN HƯỚNG DẪN")
    for _ in range(17):
        add_p(doc, "............................................................................................................................................................", line=1.3, first_line=False)
    add_p(doc, "Ngày ...... tháng ...... năm 2026", align=WD_ALIGN_PARAGRAPH.RIGHT, first_line=False, before=16)
    add_p(doc, "Giáo viên hướng dẫn", align=WD_ALIGN_PARAGRAPH.RIGHT, first_line=False)
    doc.add_page_break()

    toc(doc)
    doc.add_page_break()

    add_h1(doc, "DANH MỤC TỪ VIẾT TẮT")
    add_table(doc, ["Từ viết tắt", "Ý nghĩa"], [
        ("APK", "Android Package Kit - gói cài đặt ứng dụng Android."),
        ("CRUD", "Create, Read, Update, Delete - nhóm thao tác thêm, đọc, sửa, xóa dữ liệu."),
        ("DB", "Database - cơ sở dữ liệu."),
        ("ERD", "Entity Relationship Diagram - sơ đồ thực thể liên kết."),
        ("JDK", "Java Development Kit - bộ công cụ phát triển Java."),
        ("SDK", "Software Development Kit - bộ công cụ phát triển phần mềm."),
        ("SQLite", "Hệ quản trị cơ sở dữ liệu quan hệ nhúng trên thiết bị."),
        ("UI", "User Interface - giao diện người dùng."),
        ("UX", "User Experience - trải nghiệm người dùng."),
        ("XML", "Extensible Markup Language - ngôn ngữ đánh dấu dùng khai báo layout Android."),
    ], widths=[1.3, 5.0])
    doc.add_page_break()

    add_h1(doc, "DANH MỤC HÌNH ẢNH")
    for item in ["Hình 2.1. Kiến trúc tổng quan ứng dụng Student Planner",
                 "Hình 3.1. Mô hình dữ liệu SQLite của ứng dụng",
                 "Hình 3.2. Luồng xử lý thêm kế hoạch và nhắc lịch"]:
        add_p(doc, item, first_line=False, after=2)
    add_h1(doc, "DANH MỤC BẢNG BIỂU")
    for item in ["Bảng 1.1. Đối chiếu nội dung báo cáo với phiếu chấm",
                 "Bảng 3.1. Danh sách yêu cầu chức năng",
                 "Bảng 3.2. Danh sách yêu cầu phi chức năng",
                 "Bảng 3.3. Thiết kế các bảng dữ liệu chính",
                 "Bảng 3.4. Kịch bản kiểm thử chức năng",
                 "Bảng 4.1. Phân công và mức đóng góp cho đề tài"]:
        add_p(doc, item, first_line=False, after=2)
    doc.add_page_break()

    add_h1(doc, "TÓM TẮT ĐỀ TÀI")
    add_p(doc, "Báo cáo trình bày quá trình phân tích, thiết kế và cài đặt ứng dụng Student Planner - ứng dụng Android offline hỗ trợ sinh viên quản lý kế hoạch cá nhân. Đề tài không chỉ tập trung vào quản lý học tập mà mở rộng sang các hoạt động phổ biến trong đời sống sinh viên như làm bài tập, đi học, làm thêm, ôn thi, thực hiện dự án và các kế hoạch cá nhân.")
    add_p(doc, "Ứng dụng được xây dựng bằng Java, XML Layout, Material Components và SQLiteOpenHelper. Dữ liệu được lưu trực tiếp trong cơ sở dữ liệu SQLite nội bộ của ứng dụng trên thiết bị. Người dùng đăng ký và đăng nhập tài khoản cục bộ để tách dữ liệu giữa các sinh viên sử dụng chung thiết bị. Các chức năng chính gồm quản lý nhóm kế hoạch, quản lý kế hoạch theo loại việc, lịch ngày/tuần/tháng, nhắc lịch trước giờ hẹn, kiểm tra trùng thời gian, hiển thị kế hoạch quá hạn, thống kê tuần/tháng và gợi ý lịch hôm nay theo mức ưu tiên.")
    add_p(doc, "Kết quả đạt được là một ứng dụng có thể chạy offline, giao diện tối giản theo phong cách Material Design, dữ liệu được tổ chức phù hợp với phạm vi quản lý kế hoạch cho sinh viên. Báo cáo cũng chỉ ra các hạn chế hiện tại và hướng phát triển như đồng bộ Firebase, chia sẻ kế hoạch nhóm, sao lưu dữ liệu, tối ưu nhắc lịch và mở rộng thống kê học tập - làm thêm.")
    doc.add_page_break()

    add_h1(doc, "Chương 1:\nGIỚI THIỆU ĐỀ TÀI")
    add_h2(doc, "1.1. LÝ DO CHỌN ĐỀ TÀI")
    add_p(doc, "Sinh viên thường phải xử lý nhiều loại kế hoạch trong cùng một ngày: lịch học, hạn nộp bài tập, ca làm thêm, lịch ôn thi, đồ án nhóm và các việc cá nhân. Nếu chỉ ghi chú rời rạc hoặc chỉ quản lý theo môn học, người dùng khó quan sát tổng thể thời gian, dễ trùng lịch và dễ bỏ sót việc quan trọng. Vì vậy, đề tài lựa chọn xây dựng ứng dụng quản lý kế hoạch cho sinh viên nhằm gom các hoạt động thường ngày vào một hệ thống đơn giản, trực quan và hoạt động offline.")
    add_p(doc, "Khác với một ứng dụng ghi chú thông thường, hệ thống cần nhận biết loại kế hoạch, thời gian thực hiện, mức ưu tiên, trạng thái quá hạn, lịch lặp, nhắc lịch và thống kê theo tuần/tháng. Đây là những yếu tố làm rõ phạm vi đề tài và giúp ứng dụng có giá trị thực tiễn đối với sinh viên.")
    add_h2(doc, "1.2. MỤC TIÊU ĐỀ TÀI")
    for item in ["Xây dựng ứng dụng Android hỗ trợ sinh viên tạo, sửa, xóa và theo dõi kế hoạch cá nhân.",
                 "Phân loại kế hoạch theo nhóm và loại việc như bài tập, đi học, làm thêm, cá nhân, thi và dự án.",
                 "Lưu trữ dữ liệu cục bộ bằng SQLite, hoạt động không cần kết nối Internet.",
                 "Cung cấp lịch ngày/tuần/tháng, nhắc việc, cảnh báo trùng giờ và thống kê tiến độ.",
                 "Tự động hiển thị kế hoạch quá hạn và gợi ý việc cần ưu tiên trong ngày.",
                 "Thiết kế giao diện theo phong cách Material Design, dễ dùng trên điện thoại Android."]:
        add_bullet(doc, item)
    add_h2(doc, "1.3. ĐỐI TƯỢNG VÀ PHẠM VI NGHIÊN CỨU")
    add_p(doc, "Đối tượng sử dụng của ứng dụng là sinh viên cần quản lý lịch học, bài tập, ca làm thêm, ôn thi, dự án và các kế hoạch cá nhân. Phạm vi đề tài tập trung vào ứng dụng Android offline, không xây dựng máy chủ riêng, không đồng bộ dữ liệu đám mây và không triển khai đa nền tảng. Dữ liệu tài khoản và kế hoạch được lưu trong vùng dữ liệu riêng của ứng dụng trên thiết bị Android.")
    add_h2(doc, "1.4. PHƯƠNG PHÁP THỰC HIỆN")
    for item in ["Khảo sát nhu cầu quản lý kế hoạch của sinh viên và xác định các loại kế hoạch đặc thù.",
                 "Phân tích yêu cầu chức năng, yêu cầu phi chức năng và xây dựng mô hình dữ liệu phù hợp.",
                 "Thiết kế giao diện theo hướng tối giản, ưu tiên thao tác nhanh trên điện thoại.",
                 "Cài đặt ứng dụng Android bằng Java, XML Layout, SQLiteOpenHelper và Material Components.",
                 "Kiểm thử thủ công các luồng chính: đăng ký, đăng nhập, thêm kế hoạch, lọc, lịch, nhắc việc, thống kê và xử lý quá hạn."]:
        add_bullet(doc, item)
    add_h2(doc, "1.5. ĐỐI CHIẾU PHIẾU CHẤM")
    add_table(doc, ["Tiêu chí phiếu chấm", "Nội dung đã đáp ứng trong báo cáo"], [
        ("Hình thức báo cáo", "Có bìa, lời cảm ơn, nhận xét GVHD, mục lục, danh mục từ viết tắt, danh mục hình, danh mục bảng và tài liệu tham khảo."),
        ("Bố cục/Mở đầu", "Chương 1 trình bày lý do chọn đề tài, mục tiêu, đối tượng, phạm vi và phương pháp thực hiện."),
        ("Cơ sở lý thuyết", "Chương 2 trình bày Android, Java, XML Layout, SQLite, Material Design, notification, bảo mật mật khẩu và nghiệp vụ quản lý kế hoạch sinh viên."),
        ("Giải pháp thực hiện", "Chương 3 mô tả bài toán, yêu cầu chức năng, phi chức năng, mô hình dữ liệu, mô hình xử lý và thiết kế giao diện."),
        ("Thực nghiệm/Cài đặt/Kết quả", "Chương 3 mô tả các màn hình, chức năng CRUD, lọc, lịch, nhắc lịch, thống kê, kiểm tra trùng giờ và kiểm thử."),
        ("Kết luận và hướng phát triển", "Chương 4 tổng kết kết quả, ưu nhược điểm, hướng phát triển và mức đóng góp của thành viên."),
    ], widths=[2.25, 4.1], font_size=11.0)
    add_caption(doc, "Bảng 1.1. Đối chiếu nội dung báo cáo với phiếu chấm")
    doc.add_page_break()

    add_h1(doc, "Chương 2:\nCƠ SỞ LÝ THUYẾT")
    theory_sections = [
        ("2.1. TỔNG QUAN LẬP TRÌNH ANDROID", [
            "Android là nền tảng phổ biến cho thiết bị di động, hỗ trợ xây dựng ứng dụng bằng Java hoặc Kotlin thông qua Android SDK. Trong đề tài này, nhóm sử dụng Java để xử lý logic nghiệp vụ và XML Layout để khai báo giao diện. Cách tiếp cận này phù hợp với học phần Lập trình thiết bị di động vì giúp tách tương đối rõ phần giao diện, xử lý sự kiện và truy cập dữ liệu.",
            "Một ứng dụng Android thường được tổ chức thành các thành phần như Activity, Fragment, Adapter, BroadcastReceiver và lớp xử lý dữ liệu. Student Planner áp dụng mô hình này để giữ code dễ đọc và phù hợp với quy mô đồ án.",
        ]),
        ("2.2. SQLITE VÀ LƯU TRỮ OFFLINE", [
            "SQLite là hệ quản trị cơ sở dữ liệu quan hệ nhúng, phù hợp cho ứng dụng di động cần lưu dữ liệu cục bộ. Android cung cấp SQLiteOpenHelper để tạo, nâng cấp và truy cập cơ sở dữ liệu. Với đề tài này, SQLite được dùng để lưu tài khoản, nhóm kế hoạch và toàn bộ kế hoạch của sinh viên.",
            "Dữ liệu SQLite của ứng dụng nằm trong vùng dữ liệu riêng của ứng dụng trên thiết bị Android. Mỗi tài khoản có user_id riêng, các bảng plan_categories và tasks đều lưu user_id để tách dữ liệu giữa các tài khoản.",
        ]),
        ("2.3. MATERIAL DESIGN VÀ TRẢI NGHIỆM NGƯỜI DÙNG", [
            "Material Design cung cấp các nguyên tắc thiết kế giao diện hiện đại cho ứng dụng Android như bố cục rõ ràng, màu sắc nhất quán, khoảng trắng hợp lý, thẻ nội dung, bottom navigation, chip lọc và floating action button.",
            "Trong Student Planner, giao diện ưu tiên tính tối giản và khả năng quét nhanh thông tin. Màn hình Home hiển thị tỷ lệ hoàn thành, thống kê, kế hoạch quá hạn và gợi ý hôm nay.",
        ]),
        ("2.4. NHẮC LỊCH VÀ THÔNG BÁO TRÊN ANDROID", [
            "Nhắc lịch là chức năng quan trọng để ứng dụng quản lý kế hoạch khác biệt với ứng dụng ghi chú thông thường. Khi người dùng bật nhắc lịch, ứng dụng tính thời điểm thông báo dựa trên ngày, giờ và số phút nhắc trước.",
            "Ứng dụng đăng ký BootReceiver để khôi phục các lịch nhắc sau khi thiết bị khởi động lại. Điều này giúp hạn chế tình trạng mất nhắc lịch khi điện thoại tắt nguồn hoặc restart.",
        ]),
        ("2.5. BẢO MẬT TÀI KHOẢN OFFLINE", [
            "Mặc dù ứng dụng hoạt động offline, chức năng đăng ký và đăng nhập vẫn hợp lý vì giúp tách dữ liệu giữa nhiều người dùng trên cùng thiết bị và tạo trải nghiệm cá nhân hóa.",
            "Student Planner sử dụng PBKDF2WithHmacSHA256, salt ngẫu nhiên và số vòng lặp 120.000 để băm mật khẩu trước khi lưu vào SQLite. Ứng dụng cũng có cơ chế nâng cấp hash cũ sang PBKDF2 khi người dùng đăng nhập thành công.",
        ]),
        ("2.6. NGHIỆP VỤ QUẢN LÝ KẾ HOẠCH SINH VIÊN", [
            "Kế hoạch của sinh viên có tính đa dạng hơn một ghi chú thông thường. Một bài tập cần hạn nộp, môn liên quan và trạng thái đã nộp; một buổi học cần thời gian, phòng, môn/lớp và địa điểm; một ca làm thêm cần nơi làm, thời gian bắt đầu/kết thúc và có thể có tiền công dự kiến.",
            "Ứng dụng hiện hỗ trợ các loại kế hoạch ASSIGNMENT, CLASS, PART_TIME, PERSONAL, EXAM và PROJECT. Trạng thái gồm sắp tới, đang thực hiện, hoàn thành và đã hủy. Trạng thái quá hạn được nhận diện động dựa trên ngày hiện tại.",
        ]),
    ]
    for heading, paragraphs in theory_sections:
        add_h2(doc, heading)
        for paragraph in paragraphs:
            add_p(doc, paragraph)
    doc.add_picture(str(ASSET_DIR / "architecture.png"), width=Inches(6.25))
    add_caption(doc, "Hình 2.1. Kiến trúc tổng quan ứng dụng Student Planner")
    doc.add_page_break()

    add_h1(doc, "Chương 3:\nPHÂN TÍCH, THIẾT KẾ VÀ CÀI ĐẶT")
    add_h2(doc, "3.1. MÔ TẢ BÀI TOÁN")
    add_p(doc, "Bài toán đặt ra là xây dựng một ứng dụng Android giúp sinh viên quản lý nhiều loại kế hoạch trong đời sống hằng ngày. Người dùng có thể đăng ký tài khoản cục bộ, đăng nhập, tạo các nhóm kế hoạch, thêm kế hoạch theo loại việc, xem lịch, nhận nhắc lịch, đánh dấu hoàn thành và xem thống kê.")
    add_p(doc, "Điểm quan trọng của bài toán là không xem mọi kế hoạch như một ghi chú giống nhau. Mỗi loại kế hoạch có đặc thù riêng: bài tập có trạng thái nộp, đi học có phòng/lớp, làm thêm có nơi làm và tiền công dự kiến, dự án có nội dung liên quan.")

    add_h2(doc, "3.2. YÊU CẦU CHỨC NĂNG")
    add_table(doc, ["Mã", "Chức năng", "Mô tả"], [
        ("F01", "Đăng ký tài khoản", "Kiểm tra username, email, mật khẩu và lưu mật khẩu đã băm."),
        ("F02", "Đăng nhập/đăng xuất", "Xác thực tài khoản cục bộ và lưu phiên đăng nhập bằng SharedPreferences."),
        ("F03", "Quản lý nhóm kế hoạch", "Thêm, sửa, xóa nhóm như Học tập, Bài tập, Đi học, Làm thêm, Cá nhân."),
        ("F04", "Thêm kế hoạch", "Tạo kế hoạch với ngày, giờ, loại, ưu tiên, thời lượng, nhắc lịch, lặp lại và trường riêng theo loại."),
        ("F05", "Cập nhật/xóa kế hoạch", "Sửa thông tin, đổi trạng thái, bật/tắt nhắc lịch và xóa kế hoạch."),
        ("F06", "Lọc và tìm kiếm", "Tìm kiếm theo nội dung; lọc theo trạng thái, nhóm kế hoạch và loại kế hoạch."),
        ("F07", "Lịch ngày/tuần/tháng", "Hiển thị kế hoạch theo ngày được chọn, theo tuần hoặc theo tháng."),
        ("F08", "Nhắc lịch", "Đặt thông báo đúng giờ hoặc trước 5/15/30 phút; khôi phục sau khi khởi động lại."),
        ("F09", "Cảnh báo trùng giờ", "Kiểm tra khoảng thời gian của kế hoạch mới/sửa với các kế hoạch chưa hoàn thành."),
        ("F10", "Thống kê và gợi ý", "Dashboard thống kê tổng quan, tuần/tháng và gợi ý việc hôm nay theo ưu tiên."),
    ], widths=[0.6, 1.45, 4.3], font_size=10.5)
    add_caption(doc, "Bảng 3.1. Danh sách yêu cầu chức năng")

    doc.add_page_break()
    add_h2(doc, "3.3. YÊU CẦU PHI CHỨC NĂNG")
    add_table(doc, ["Nhóm yêu cầu", "Mô tả"], [
        ("Khả dụng", "Ứng dụng phải sử dụng được khi không có Internet; các luồng chính thao tác trong vài bước và có thông báo phản hồi."),
        ("Dễ dùng", "Giao diện rõ ràng, có bottom navigation, card, chip lọc, khoảng trắng hợp lý và form động theo loại kế hoạch."),
        ("Bảo mật", "Mật khẩu được băm bằng PBKDF2 có salt; dữ liệu tài khoản và kế hoạch tách theo user_id; backup ứng dụng bị tắt trong manifest."),
        ("Bảo trì", "Mã nguồn chia theo package activity, fragment, adapter, data/model, data/local, notification và utils."),
        ("Hiệu năng", "Truy vấn dữ liệu chạy trên ExecutorService; cơ sở dữ liệu có index phục vụ lọc và thống kê."),
        ("Tương thích", "Ứng dụng hỗ trợ minSdk 23, targetSdk 35, Java 17, AppCompat, RecyclerView và Material Components."),
    ], widths=[1.55, 4.8], font_size=11)
    add_caption(doc, "Bảng 3.2. Danh sách yêu cầu phi chức năng")

    add_h2(doc, "3.4. MÔ HÌNH DỮ LIỆU")
    add_p(doc, "Cơ sở dữ liệu của ứng dụng có tên personal_planner.db, được quản lý trong lớp DatabaseHelper. Phiên bản hiện tại là version 4. Ba bảng chính gồm users, plan_categories và tasks.")
    doc.add_picture(str(ASSET_DIR / "erd.png"), width=Inches(5.15))
    add_caption(doc, "Hình 3.1. Mô hình dữ liệu SQLite của ứng dụng")
    doc.add_page_break()
    add_table(doc, ["Bảng", "Thuộc tính chính", "Khóa/quan hệ", "Vai trò"], [
        ("users", "user_id, username, email, password, created_at", "PK: user_id; username/email unique", "Lưu tài khoản cục bộ và thông tin đăng nhập đã băm mật khẩu."),
        ("plan_categories", "category_id, category_name, category_code, note, color, user_id", "PK: category_id; FK: user_id", "Lưu nhóm kế hoạch để phân loại lịch học, bài tập, làm thêm, cá nhân."),
        ("tasks", "task_id, title, description, date, time, end_time, status, category_id, plan_type, priority, duration_minutes, reminder_enabled, reminder_minutes, location, room, subject, repeat_rule, repeat_until, wage, submitted, user_id", "PK: task_id; FK: category_id, user_id", "Lưu nội dung kế hoạch, loại việc, thời gian, trạng thái, nhắc lịch và các trường riêng theo nghiệp vụ sinh viên."),
    ], widths=[1.15, 2.25, 1.35, 1.75], font_size=9.8)
    add_caption(doc, "Bảng 3.3. Thiết kế các bảng dữ liệu chính")

    add_h2(doc, "3.5. MÔ HÌNH XỬ LÝ")
    add_p(doc, "Luồng xử lý chính của ứng dụng bắt đầu từ đăng nhập. Khi đăng nhập thành công, SessionManager lưu user_id và username, MainActivity mở các tab Home, Kế hoạch, Lịch và Cá nhân. Mọi truy vấn dữ liệu đều truyền user_id để chỉ lấy dữ liệu thuộc tài khoản hiện tại.")
    add_p(doc, "Khi thêm kế hoạch, hệ thống kiểm tra tiêu đề, thời lượng và danh sách nhóm kế hoạch. Nếu kế hoạch có lặp lại, hàm buildRepeatDates tạo danh sách ngày cần lưu. Trước mỗi lần lưu, hệ thống gọi hasTimeConflict để phát hiện trùng giờ.")
    doc.add_picture(str(ASSET_DIR / "workflow.png"), width=Inches(5.35))
    add_caption(doc, "Hình 3.2. Luồng xử lý thêm kế hoạch và nhắc lịch")

    add_h2(doc, "3.6. THIẾT KẾ GIAO DIỆN VÀ KẾT QUẢ CÀI ĐẶT")
    add_p(doc, "Giao diện ứng dụng được xây dựng bằng XML Layout và Material Components. Màu chủ đạo gồm xanh dương, xanh ngọc, trắng và xám nhạt. Các khối nội dung dùng MaterialCardView, bo góc mềm, có khoảng trắng để phù hợp phong cách dashboard hiện đại.")
    for title, body in [
        ("3.6.1. Màn hình đăng ký và đăng nhập", "Màn hình đăng ký kiểm tra tên đăng nhập, email, mật khẩu và xác nhận mật khẩu. Màn hình đăng nhập xác thực tài khoản trong SQLite, xác thực mật khẩu bằng PasswordUtils và mở MainActivity nếu thành công."),
        ("3.6.2. Màn hình Trang chủ", "Trang chủ đóng vai trò dashboard. Người dùng nhìn thấy tỷ lệ hoàn thành, tổng số kế hoạch, số bài tập, giờ đi học, giờ làm thêm, số kế hoạch quá hạn, thống kê tuần/tháng và danh sách gợi ý hôm nay."),
        ("3.6.3. Màn hình danh sách kế hoạch", "Màn hình danh sách dùng RecyclerView để hiển thị kế hoạch. Người dùng có thể tìm kiếm theo từ khóa, lọc theo trạng thái, nhóm kế hoạch và loại kế hoạch. Kế hoạch quá hạn được hiển thị bằng badge, màu đỏ ở ngày giờ và viền card đỏ."),
        ("3.6.4. Màn hình thêm/sửa kế hoạch", "Form thêm/sửa kế hoạch thay đổi trường nhập theo loại kế hoạch. Với bài tập, ứng dụng hiển thị môn liên quan và trạng thái đã nộp. Với đi học, ứng dụng hiển thị địa điểm, phòng/ca học và môn/lớp. Với làm thêm, ứng dụng hiển thị nơi làm và tiền công dự kiến."),
        ("3.6.5. Màn hình lịch", "Màn hình lịch dùng CalendarView kết hợp chip Ngày/Tuần/Tháng. Khi chọn ngày hoặc chế độ xem, ứng dụng truy vấn kế hoạch đúng khoảng thời gian và hiển thị trong RecyclerView."),
    ]:
        add_h3(doc, title)
        add_p(doc, body)

    add_h2(doc, "3.7. CÀI ĐẶT CÁC CHỨC NĂNG CHÍNH")
    add_p(doc, "Về lưu trữ, DatabaseHelper kế thừa SQLiteOpenHelper, tạo bảng users, plan_categories và tasks. Các phương thức addStudyPlan, updateStudyPlan, deleteStudyPlan, getStudyPlans, getStudyPlansByDate, getStudyPlansBetween, getUpcomingPlans, getTodaySuggestions và getPlanRangeStats phục vụ các màn hình chính.")
    add_p(doc, "Về nhắc lịch, ReminderScheduler dùng AlarmManager.setAndAllowWhileIdle để đặt thời điểm thông báo. ReminderReceiver tạo notification với kênh study_reminders và mở MainActivity khi người dùng chạm vào thông báo. BootReceiver lắng nghe BOOT_COMPLETED để lên lịch lại các kế hoạch có reminder_enabled = 1.")
    add_p(doc, "Về thống kê, StudyStatistics phục vụ thống kê tổng quan toàn bộ dữ liệu, còn PlanRangeStats phục vụ thống kê theo tuần/tháng. Dashboard hiện giờ học, giờ làm thêm, bài tập chưa nộp, tỷ lệ hoàn thành và số kế hoạch quá hạn.")

    add_h2(doc, "3.8. KIỂM THỬ CHỨC NĂNG")
    add_table(doc, ["Mã", "Kịch bản kiểm thử", "Kết quả mong đợi", "Kết quả"], [
        ("TC01", "Đăng ký tài khoản mới.", "Tạo tài khoản thành công, mật khẩu được lưu dạng hash.", "Đạt"),
        ("TC02", "Đăng nhập bằng tài khoản đã đăng ký.", "Mở MainActivity và lưu phiên đăng nhập.", "Đạt"),
        ("TC03", "Thêm kế hoạch bài tập có ưu tiên cao và bật nhắc trước 15 phút.", "Kế hoạch được lưu và hiển thị trong danh sách.", "Đạt"),
        ("TC04", "Thêm kế hoạch đi học trùng giờ.", "Ứng dụng cảnh báo trùng giờ hoặc bỏ qua ngày trùng khi tạo lặp.", "Đạt"),
        ("TC05", "Lọc danh sách theo loại Làm thêm.", "Chỉ hiển thị kế hoạch PART_TIME.", "Đạt"),
        ("TC06", "Chọn chế độ lịch Tuần.", "Hiển thị kế hoạch trong khoảng thứ Hai đến Chủ nhật.", "Đạt"),
        ("TC07", "Tạo kế hoạch ngày trước hôm nay và chưa hoàn thành.", "Hiển thị badge Quá hạn, ngày giờ màu đỏ và viền card đỏ.", "Đạt"),
        ("TC08", "Đánh dấu kế hoạch hoàn thành.", "Trạng thái cập nhật và reminder bị hủy nếu có.", "Đạt"),
        ("TC09", "Mở Home sau khi có nhiều loại kế hoạch.", "Dashboard hiển thị thống kê, tuần/tháng và gợi ý hôm nay.", "Đạt"),
    ], widths=[0.55, 2.35, 2.55, 0.85], font_size=10.0)
    add_caption(doc, "Bảng 3.4. Kịch bản kiểm thử chức năng")
    add_p(doc, "Ngoài kiểm thử thủ công, nhóm đã biên dịch ứng dụng bằng Gradle với tác vụ assembleDebug. Kết quả build thành công cho thấy mã nguồn, tài nguyên XML, manifest và dependencies tương thích ở thời điểm kiểm tra.")
    doc.add_page_break()

    add_h1(doc, "Chương 4:\nKẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN")
    add_h2(doc, "4.1. KẾT QUẢ ĐẠT ĐƯỢC")
    add_p(doc, "Sau quá trình thực hiện, nhóm đã xây dựng được ứng dụng Android quản lý kế hoạch cho sinh viên theo đúng phạm vi đề tài. Ứng dụng có đăng ký/đăng nhập offline, quản lý nhóm kế hoạch, quản lý kế hoạch theo loại việc, form động, lịch ngày/tuần/tháng, nhắc lịch, kiểm tra trùng giờ, thống kê tổng quan, thống kê tuần/tháng, gợi ý lịch hôm nay và nhận diện kế hoạch quá hạn.")
    add_p(doc, "Về mặt kỹ thuật, ứng dụng sử dụng Java, XML Layout, Material Components, SQLiteOpenHelper, SharedPreferences, AlarmManager, BroadcastReceiver và NotificationCompat. Cấu trúc mã nguồn được chia thành các package rõ ràng giúp dễ bảo trì.")
    add_h2(doc, "4.2. ƯU ĐIỂM VÀ NHƯỢC ĐIỂM")
    add_h3(doc, "4.2.1. Ưu điểm")
    for item in ["Ứng dụng hoạt động offline, không phụ thuộc mạng Internet.",
                 "Mô hình dữ liệu phản ánh đúng nhiều loại kế hoạch của sinh viên.",
                 "Giao diện dashboard hiện đại, có thống kê và gợi ý hôm nay.",
                 "Danh sách kế hoạch hỗ trợ tìm kiếm, lọc theo trạng thái, nhóm và loại kế hoạch.",
                 "Có nhắc lịch trước giờ hẹn, khôi phục nhắc lịch sau khi khởi động thiết bị.",
                 "Mật khẩu được băm bằng PBKDF2 có salt, tốt hơn lưu plain text hoặc SHA-256 đơn giản."]:
        add_bullet(doc, item)
    add_h3(doc, "4.2.2. Nhược điểm")
    for item in ["Chưa đồng bộ dữ liệu giữa nhiều thiết bị vì ứng dụng ưu tiên offline.",
                 "Chưa có màn hình chi tiết riêng cho từng loại kế hoạch với bố cục chuyên sâu hơn.",
                 "Thông báo hiện mở về MainActivity, chưa mở thẳng vào chi tiết kế hoạch.",
                 "Chưa có biểu đồ thống kê trực quan theo tuần/tháng.",
                 "Chưa hỗ trợ chia sẻ kế hoạch nhóm hoặc cộng tác giữa nhiều sinh viên."]:
        add_bullet(doc, item)
    add_h2(doc, "4.3. HƯỚNG PHÁT TRIỂN")
    for item in ["Tích hợp Firebase Authentication và Cloud Firestore để đồng bộ dữ liệu, sao lưu và đăng nhập đa thiết bị.",
                 "Bổ sung biểu đồ thống kê theo tuần/tháng: giờ học, giờ làm thêm, số bài tập chưa nộp, tỷ lệ hoàn thành và số kế hoạch quá hạn.",
                 "Mở rộng nhắc lịch: nhắc lại nhiều lần, tùy chỉnh âm thanh, mở thẳng vào màn hình chi tiết kế hoạch.",
                 "Bổ sung chức năng xuất/nhập dữ liệu hoặc sao lưu file để tránh mất dữ liệu khi đổi thiết bị.",
                 "Xây dựng chức năng chia sẻ kế hoạch nhóm cho đồ án, học nhóm hoặc lịch làm thêm theo ca.",
                 "Tối ưu giao diện cho tablet và hỗ trợ dark mode đầy đủ hơn."]:
        add_bullet(doc, item)
    add_h2(doc, "4.4. PHÂN CÔNG VÀ MỨC ĐÓNG GÓP")
    add_table(doc, ["Thành viên", "Nhiệm vụ chính", "Mức đóng góp"], [
        ("SV1 - MSSV", "Phân tích yêu cầu, thiết kế cơ sở dữ liệu, cài đặt DatabaseHelper, đăng ký/đăng nhập và bảo mật mật khẩu.", "33%"),
        ("SV2 - MSSV", "Thiết kế giao diện, xây dựng Home/Task/Calendar/Profile, lọc kế hoạch, thống kê và hiển thị quá hạn.", "33%"),
        ("SV3 - MSSV", "Cài đặt thêm/sửa/xóa kế hoạch, nhắc lịch, lặp lịch, kiểm tra trùng giờ, kiểm thử và hoàn thiện báo cáo.", "34%"),
    ], widths=[1.35, 4.25, 0.8], font_size=10.8)
    add_caption(doc, "Bảng 4.1. Phân công và mức đóng góp cho đề tài")
    doc.add_page_break()

    add_h1(doc, "TÀI LIỆU THAM KHẢO")
    refs = [
        "[1] Android Developers, “Guide to app architecture”, https://developer.android.com/topic/architecture, truy cập ngày 22/06/2026.",
        "[2] Android Developers, “Save data using SQLite”, https://developer.android.com/training/data-storage/sqlite, truy cập ngày 22/06/2026.",
        "[3] Android Developers, “About notifications in Views”, https://developer.android.com/develop/ui/views/notifications, truy cập ngày 22/06/2026.",
        "[4] Material Design, “Material Design 3”, https://m3.material.io/, truy cập ngày 22/06/2026.",
        "[5] OWASP Foundation, “Password Storage Cheat Sheet”, https://cheatsheetseries.owasp.org/cheatsheets/Password_Storage_Cheat_Sheet.html, truy cập ngày 22/06/2026.",
        "[6] Oracle, “Java Platform, Standard Edition Documentation”, https://docs.oracle.com/en/java/javase/, truy cập ngày 22/06/2026.",
        "[7] Mã nguồn ứng dụng Student Planner trong thư mục D:\\LTDD\\android_todoapp-main, truy cập và kiểm tra ngày 22/06/2026.",
    ]
    for ref in refs:
        add_p(doc, ref, first_line=False, after=4, line=1.25)

    set_footer(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
