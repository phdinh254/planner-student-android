from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"D:\LTDD\PersonalPlanner_Java_XML_SQLite\PersonalPlanner")
ASSET_DIR = ROOT / "report_assets"
OUTPUT = ROOT / "Bao_cao_do_an_Study_Planner.docx"
ASSET_DIR.mkdir(exist_ok=True)

FONT_REGULAR = Path(r"C:\Windows\Fonts\arial.ttf")
FONT_BOLD = Path(r"C:\Windows\Fonts\arialbd.ttf")
TITLE = "XÂY DỰNG ỨNG DỤNG QUẢN LÝ KẾ HOẠCH HỌC TẬP CÁ NHÂN DÀNH CHO SINH VIÊN TRÊN NỀN TẢNG ANDROID"
SHORT_TITLE = "Ứng dụng quản lý kế hoạch học tập cá nhân dành cho sinh viên"
STUDENTS = "Trần Phú Dinh, Sơn Trung Nguyên, Lê Đăng Huy"


def font(size, bold=False):
    path = FONT_BOLD if bold else FONT_REGULAR
    return ImageFont.truetype(str(path), size=size)


def draw_wrapped(draw, xy, text, max_chars, used_font, fill="#17201E",
                 anchor="mm", spacing=6, align="center"):
    lines = []
    for paragraph in text.split("\n"):
        lines.extend(wrap(paragraph, max_chars) or [""])
    draw.multiline_text(xy, "\n".join(lines), font=used_font, fill=fill,
                        anchor=anchor, spacing=spacing, align=align)


def rounded(draw, box, radius=24, fill="#FFFFFF", outline="#CBD4D0", width=3):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw, start, end, fill="#1F6F68", width=5):
    draw.line([start, end], fill=fill, width=width)
    x1, y1 = start
    x2, y2 = end
    dx, dy = x2 - x1, y2 - y1
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    size = 16
    p1 = (x2, y2)
    p2 = (x2 - ux * size + px * size * 0.55, y2 - uy * size + py * size * 0.55)
    p3 = (x2 - ux * size - px * size * 0.55, y2 - uy * size - py * size * 0.55)
    draw.polygon([p1, p2, p3], fill=fill)


def make_use_case(path):
    img = Image.new("RGB", (1600, 1000), "#F6F7F4")
    d = ImageDraw.Draw(img)
    rounded(d, (250, 55, 1540, 940), 26, "#FFFFFF", "#1F6F68", 4)
    d.text((895, 90), "HỆ THỐNG STUDY PLANNER", font=font(34, True),
           fill="#15534E", anchor="mm")

    # Actor
    d.ellipse((70, 280, 150, 360), outline="#17201E", width=5)
    d.line((110, 360, 110, 510), fill="#17201E", width=5)
    d.line((45, 420, 175, 420), fill="#17201E", width=5)
    d.line((110, 510, 55, 610), fill="#17201E", width=5)
    d.line((110, 510, 165, 610), fill="#17201E", width=5)
    d.text((110, 655), "Sinh viên", font=font(28, True), fill="#17201E", anchor="mm")

    cases = [
        (470, 200, "Đăng ký / đăng nhập"),
        (830, 200, "Quản lý môn học"),
        (1190, 200, "Quản lý kế hoạch học tập"),
        (470, 420, "Tìm kiếm và lọc"),
        (830, 420, "Xem lịch học"),
        (1190, 420, "Đánh dấu hoàn thành"),
        (470, 650, "Nhận thông báo nhắc học"),
        (830, 650, "Xem thống kê tiến độ"),
        (1190, 650, "Quản lý hồ sơ / đăng xuất"),
    ]
    for x, y, text in cases:
        d.ellipse((x - 160, y - 65, x + 160, y + 65),
                  fill="#D5ECE8", outline="#1F6F68", width=3)
        draw_wrapped(d, (x, y), text, 24, font(23, True))
        arrow(d, (180, 420), (x - 165, y), "#5D6865", 3)

    img.save(path, quality=95)


def make_erd(path):
    img = Image.new("RGB", (1600, 980), "#F6F7F4")
    d = ImageDraw.Draw(img)
    d.text((800, 50), "SƠ ĐỒ QUAN HỆ DỮ LIỆU (ERD)", font=font(36, True),
           fill="#15534E", anchor="mm")

    entities = {
        "USERS": (70, 190, 480, 720, [
            "PK user_id: INTEGER",
            "username: TEXT UNIQUE",
            "email: TEXT UNIQUE",
            "password: TEXT",
            "created_at: TEXT",
        ]),
        "COURSES": (595, 115, 1030, 650, [
            "PK course_id: INTEGER",
            "course_name: TEXT",
            "course_code: TEXT",
            "lecturer: TEXT",
            "color: TEXT",
            "FK user_id: INTEGER",
        ]),
        "STUDY_PLANS (tasks)": (1120, 110, 1570, 870, [
            "PK task_id: INTEGER",
            "title: TEXT",
            "description: TEXT",
            "date, time: TEXT",
            "status: INTEGER",
            "FK course_id: INTEGER",
            "priority: INTEGER",
            "duration_minutes: INTEGER",
            "reminder_enabled: INTEGER",
            "FK user_id: INTEGER",
        ]),
    }
    for name, (x1, y1, x2, y2, fields) in entities.items():
        rounded(d, (x1, y1, x2, y2), 22, "#FFFFFF", "#1F6F68", 4)
        d.rounded_rectangle((x1, y1, x2, y1 + 85), radius=22, fill="#1F6F68")
        d.rectangle((x1, y1 + 55, x2, y1 + 85), fill="#1F6F68")
        d.text(((x1 + x2) / 2, y1 + 42), name, font=font(28, True),
               fill="#FFFFFF", anchor="mm")
        y = y1 + 115
        for field in fields:
            d.text((x1 + 25, y), field, font=font(22), fill="#17201E")
            y += 55

    arrow(d, (480, 330), (595, 300), "#D08A4B", 5)
    d.text((540, 270), "1 : N", font=font(22, True), fill="#A65F28", anchor="mm")
    arrow(d, (480, 560), (1120, 700), "#D08A4B", 5)
    d.text((800, 660), "1 : N", font=font(22, True), fill="#A65F28", anchor="mm")
    arrow(d, (1030, 360), (1120, 340), "#D08A4B", 5)
    d.text((1075, 300), "1 : N", font=font(22, True), fill="#A65F28", anchor="mm")
    img.save(path, quality=95)


def make_architecture(path):
    img = Image.new("RGB", (1600, 950), "#F6F7F4")
    d = ImageDraw.Draw(img)
    d.text((800, 55), "KIẾN TRÚC ỨNG DỤNG", font=font(38, True),
           fill="#15534E", anchor="mm")

    layers = [
        (120, 140, 1480, 310, "#D5ECE8", "TẦNG GIAO DIỆN",
         "Activity · Fragment · XML Layout · Material Components · RecyclerView Adapter"),
        (120, 385, 1480, 555, "#FFF1E4", "TẦNG NGHIỆP VỤ VÀ ĐIỀU PHỐI",
         "Validation · SessionManager · tìm kiếm/lọc · thống kê · lập lịch nhắc học"),
        (120, 630, 1480, 800, "#E9EEF6", "TẦNG DỮ LIỆU VÀ HỆ THỐNG",
         "DatabaseHelper · SQLite · Model · SharedPreferences · AlarmManager · BroadcastReceiver"),
    ]
    for x1, y1, x2, y2, color, title, desc in layers:
        rounded(d, (x1, y1, x2, y2), 28, color, "#1F6F68", 3)
        d.text((800, y1 + 50), title, font=font(28, True), fill="#15534E", anchor="mm")
        d.text((800, y1 + 112), desc, font=font(23), fill="#17201E", anchor="mm")
    arrow(d, (800, 310), (800, 385), "#1F6F68", 7)
    arrow(d, (800, 555), (800, 630), "#1F6F68", 7)
    d.text((800, 875), "Dữ liệu được phân tách theo từng user_id; thao tác SQLite chạy trên ExecutorService.",
           font=font(23, True), fill="#5D6865", anchor="mm")
    img.save(path, quality=95)


def make_flow(path):
    img = Image.new("RGB", (1600, 950), "#F6F7F4")
    d = ImageDraw.Draw(img)
    d.text((800, 50), "LUỒNG HOẠT ĐỘNG CHÍNH", font=font(38, True),
           fill="#15534E", anchor="mm")
    nodes = [
        (140, 150, 430, 260, "Mở ứng dụng"),
        (650, 150, 950, 260, "Kiểm tra phiên đăng nhập"),
        (1170, 150, 1470, 260, "Đăng nhập / đăng ký"),
        (650, 370, 950, 480, "Màn hình chính"),
        (100, 650, 380, 770, "Tổng quan\nthống kê"),
        (440, 650, 720, 770, "Danh sách\nkế hoạch"),
        (780, 650, 1060, 770, "Lịch học"),
        (1120, 650, 1500, 770, "Hồ sơ và\nquản lý môn học"),
    ]
    for x1, y1, x2, y2, text in nodes:
        rounded(d, (x1, y1, x2, y2), 25, "#FFFFFF", "#1F6F68", 4)
        draw_wrapped(d, ((x1 + x2) / 2, (y1 + y2) / 2), text, 25, font(24, True))
    arrow(d, (430, 205), (650, 205))
    arrow(d, (950, 205), (1170, 205))
    d.text((1060, 170), "Chưa có phiên", font=font(19, True), fill="#A65F28", anchor="mm")
    arrow(d, (800, 260), (800, 370))
    d.text((850, 320), "Đã đăng nhập", font=font(19, True), fill="#15534E")
    arrow(d, (1320, 260), (950, 425))
    for x in [240, 580, 920, 1310]:
        arrow(d, (800, 480), (x, 650), "#5D6865", 4)
    img.save(path, quality=95)


def phone_frame(draw, box, title, accent="#1F6F68"):
    x1, y1, x2, y2 = box
    rounded(draw, box, 42, "#FFFFFF", "#17201E", 5)
    draw.rounded_rectangle((x1 + 85, y1 + 16, x2 - 85, y1 + 32), radius=8, fill="#17201E")
    draw.text((x1 + 35, y1 + 65), title, font=font(22, True), fill=accent)
    return x1, y1, x2, y2


def make_ui_mockups(path):
    img = Image.new("RGB", (1800, 1180), "#EEF2EF")
    d = ImageDraw.Draw(img)
    d.text((900, 45), "MÔ PHỎNG BỐ CỤC CÁC MÀN HÌNH CHÍNH",
           font=font(38, True), fill="#15534E", anchor="mm")
    boxes = [(70, 110, 540, 1080), (665, 110, 1135, 1080), (1260, 110, 1730, 1080)]

    # Login
    x1, y1, x2, y2 = phone_frame(d, boxes[0], "Study Planner")
    d.text(((x1 + x2) / 2, y1 + 155), "Quản lý môn học, lịch học\nvà tiến độ cá nhân.",
           font=font(22), fill="#5D6865", anchor="mm", align="center")
    for y, label in [(300, "Tên đăng nhập"), (430, "Mật khẩu")]:
        rounded(d, (x1 + 35, y1 + y, x2 - 35, y1 + y + 85), 18, "#FFFFFF", "#CBD4D0", 3)
        d.text((x1 + 60, y1 + y + 42), label, font=font(20), fill="#5D6865", anchor="lm")
    rounded(d, (x1 + 35, y1 + 570, x2 - 35, y1 + 655), 18, "#1F6F68", "#1F6F68", 2)
    d.text(((x1 + x2) / 2, y1 + 612), "Đăng nhập", font=font(23, True),
           fill="#FFFFFF", anchor="mm")
    d.text(((x1 + x2) / 2, y1 + 720), "Chưa có tài khoản? Đăng ký",
           font=font(19, True), fill="#1F6F68", anchor="mm")

    # Home
    x1, y1, x2, y2 = phone_frame(d, boxes[1], "Xin chào, sinh viên")
    d.text((x1 + 35, y1 + 110), "Theo dõi khối lượng học tập", font=font(19),
           fill="#5D6865")
    rounded(d, (x1 + 30, y1 + 170, x2 - 30, y1 + 350), 22, "#D5ECE8", "#D5ECE8", 1)
    d.text((x1 + 55, y1 + 205), "Tiến độ học tập", font=font(23, True), fill="#15534E")
    d.text((x1 + 55, y1 + 255), "Bạn đã hoàn thành 67% kế hoạch.", font=font(18), fill="#5D6865")
    d.rounded_rectangle((x1 + 55, y1 + 305, x2 - 55, y1 + 325), radius=10, fill="#FFFFFF")
    d.rounded_rectangle((x1 + 55, y1 + 305, x1 + 295, y1 + 325), radius=10, fill="#1F6F68")
    stats = [("8", "Tổng"), ("3", "Chưa xong"), ("5", "Hoàn thành"), ("4", "Môn học")]
    for index, (value, label) in enumerate(stats):
        col = index % 2
        row = index // 2
        bx1 = x1 + 30 + col * 205
        by1 = y1 + 385 + row * 175
        rounded(d, (bx1, by1, bx1 + 190, by1 + 145), 20, "#FFFFFF", "#CBD4D0", 2)
        d.text((bx1 + 25, by1 + 25), value, font=font(30, True), fill="#1F6F68")
        d.text((bx1 + 25, by1 + 90), label, font=font(17), fill="#5D6865")
    d.text(((x1 + x2) / 2, y2 - 65), "Tổng quan   Kế hoạch   +   Lịch   Cá nhân",
           font=font(15, True), fill="#1F6F68", anchor="mm")

    # Plan list
    x1, y1, x2, y2 = phone_frame(d, boxes[2], "Kế hoạch học tập")
    rounded(d, (x1 + 30, y1 + 115, x2 - 30, y1 + 190), 18, "#FFFFFF", "#CBD4D0", 2)
    d.text((x1 + 55, y1 + 153), "Tìm theo nội dung hoặc môn học",
           font=font(17), fill="#5D6865", anchor="lm")
    d.text((x1 + 35, y1 + 225), "Tất cả     Chưa xong     Hoàn thành",
           font=font(16, True), fill="#1F6F68")
    cards = [
        ("Lập trình Android", "Hoàn thành giao diện đăng nhập", "Cao · 90 phút", "#3568A8"),
        ("Cơ sở dữ liệu", "Ôn tập truy vấn SQLite", "Trung bình · 60 phút", "#1F6F68"),
        ("Mạng máy tính", "Đọc chương mô hình OSI", "Thấp · 45 phút", "#8A5A9E"),
    ]
    for index, (course, title, meta, color) in enumerate(cards):
        top = y1 + 285 + index * 205
        rounded(d, (x1 + 30, top, x2 - 30, top + 175), 20, "#FFFFFF", "#CBD4D0", 2)
        d.rectangle((x1 + 30, top, x1 + 42, top + 175), fill=color)
        d.text((x1 + 65, top + 25), course, font=font(17, True), fill=color)
        d.text((x1 + 65, top + 65), title, font=font(20, True), fill="#17201E")
        d.text((x1 + 65, top + 110), meta, font=font(16), fill="#5D6865")
        d.text((x1 + 65, top + 140), "2026-06-15 · 19:30", font=font(15, True), fill="#1F6F68")
    img.save(path, quality=95)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=11):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, color="FFFFFF")
        set_cell_shading(table.rows[0].cells[i], "1F6F68")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if len(table.rows) % 2 == 1:
                set_cell_shading(cells[i], "F2F6F4")
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    return table


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_field(paragraph, "PAGE")


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    if bold_lead and text.startswith(bold_lead):
        p.add_run(bold_lead).bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.add_run(item)


def add_caption(doc, text):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(text).bold = True
    return p


def add_figure(doc, path, caption, width=6.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    for name, size, color in [
        ("Title", 20, "15534E"),
        ("Heading 1", 16, "15534E"),
        ("Heading 2", 14, "1F6F68"),
        ("Heading 3", 13, "17201E"),
    ]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
    doc.styles["Heading 1"].paragraph_format.page_break_before = True
    doc.styles["Heading 1"].paragraph_format.space_after = Pt(12)
    doc.styles["Heading 2"].paragraph_format.space_before = Pt(10)
    doc.styles["Heading 2"].paragraph_format.space_after = Pt(6)
    caption = doc.styles["Caption"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    caption.font.size = Pt(11)
    caption.font.bold = True
    caption.font.color.rgb = RGBColor.from_string("17201E")


def configure_section(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)


def set_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(SHORT_TITLE)
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(93, 104, 101)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(STUDENTS + "  |  Trang ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)
    add_field(p, "PAGE")


def make_document():
    use_case = ASSET_DIR / "use_case.png"
    erd = ASSET_DIR / "erd.png"
    architecture = ASSET_DIR / "architecture.png"
    flow = ASSET_DIR / "flow.png"
    ui_mockups = ASSET_DIR / "ui_mockups.png"
    make_use_case(use_case)
    make_erd(erd)
    make_architecture(architecture)
    make_flow(flow)
    make_ui_mockups(ui_mockups)

    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    doc.sections[0].different_first_page_header_footer = True
    set_header_footer(doc.sections[0])
    props = doc.core_properties
    props.title = TITLE
    props.subject = "Báo cáo đồ án môn Lập trình thiết bị di động"
    props.author = STUDENTS
    props.keywords = "Android, Java, XML, SQLite, Study Planner, AlarmManager"

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("TRƯỜNG KỸ THUẬT VÀ CÔNG NGHỆ")
    r.bold = True
    r.font.size = Pt(15)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("KHOA CÔNG NGHỆ THÔNG TIN")
    r.bold = True
    r.font.size = Pt(15)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    r = p.add_run("BÁO CÁO ĐỒ ÁN MÔN HỌC")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(21, 83, 78)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("LẬP TRÌNH THIẾT BỊ DI ĐỘNG")
    r.bold = True
    r.font.size = Pt(18)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(42)
    r = p.add_run("ĐỀ TÀI")
    r.bold = True
    r.font.size = Pt(16)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run(TITLE)
    r.bold = True
    r.font.size = Pt(19)
    r.font.color.rgb = RGBColor(31, 111, 104)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(55)
    p.paragraph_format.left_indent = Cm(3.5)
    r = p.add_run("Giảng viên hướng dẫn: ")
    r.bold = True
    p.add_run("Thầy Trịnh Quốc Việt")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(3.5)
    r = p.add_run("Sinh viên thực hiện:")
    r.bold = True
    for student in ["Trần Phú Dinh", "Sơn Trung Nguyên", "Lê Đăng Huy"]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(6)
        p.add_run(student)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(55)
    r = p.add_run("Vĩnh Long, năm 2026")
    r.bold = True
    page_break(doc)

    # Review pages
    for heading, signer in [
        ("NHẬN XÉT CỦA GIẢNG VIÊN HƯỚNG DẪN", "Giảng viên hướng dẫn"),
        ("NHẬN XÉT CỦA THÀNH VIÊN HỘI ĐỒNG", "Thành viên hội đồng"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(heading)
        r.bold = True
        r.font.size = Pt(16)
        for _ in range(17):
            doc.add_paragraph("............................................................................................................................")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run("Vĩnh Long, ngày …… tháng …… năm 2026\n")
        p.add_run(signer + "\n(Ký tên và ghi rõ họ tên)")
        page_break(doc)

    doc.add_heading("LỜI CẢM ƠN", level=1)
    add_body(doc, "Nhóm chúng em xin trân trọng gửi lời cảm ơn đến Giảng viên hướng dẫn – Thầy Trịnh Quốc Việt – đã tận tình định hướng, góp ý chuyên môn và hỗ trợ trong quá trình thực hiện đồ án môn Lập trình thiết bị di động với đề tài “Xây dựng ứng dụng quản lý kế hoạch học tập cá nhân dành cho sinh viên trên nền tảng Android”. Những góp ý của Thầy giúp nhóm hiểu rõ hơn quy trình phát triển một ứng dụng Android hoàn chỉnh, từ khảo sát yêu cầu, thiết kế dữ liệu, xây dựng giao diện đến kiểm thử và hoàn thiện sản phẩm.")
    add_body(doc, "Nhóm cũng xin cảm ơn Khoa Công nghệ Thông tin – Trường Kỹ thuật và Công nghệ – đã tạo điều kiện về môi trường học tập, tài liệu và cơ sở thực hành. Kiến thức từ các học phần lập trình Java, cơ sở dữ liệu và lập trình thiết bị di động là nền tảng quan trọng để nhóm hiện thực hóa ứng dụng Study Planner.")
    add_body(doc, "Mặc dù nhóm đã cố gắng kiểm tra và hoàn thiện báo cáo, sản phẩm vẫn có thể còn hạn chế do thời gian và kinh nghiệm thực tế. Nhóm mong nhận được ý kiến đóng góp của quý Thầy, Cô để tiếp tục cải thiện kỹ năng phân tích, thiết kế và phát triển phần mềm.")
    add_body(doc, "Nhóm chúng em xin chân thành cảm ơn!")

    doc.add_heading("LỜI CAM ĐOAN", level=1)
    add_body(doc, "Nhóm cam đoan nội dung báo cáo và sản phẩm Study Planner là kết quả nghiên cứu, thiết kế và lập trình của nhóm dưới sự hướng dẫn của giảng viên. Các tài liệu tham khảo được ghi rõ nguồn ở cuối báo cáo. Nhóm chịu trách nhiệm về tính trung thực của nội dung, số liệu kiểm thử và kết quả trình bày.")

    doc.add_heading("MỤC LỤC", level=1)
    toc = doc.add_paragraph()
    add_field(toc, 'TOC \\o "1-3" \\h \\z \\u')
    add_body(doc, "Lưu ý: Mục lục được cập nhật tự động khi mở tài liệu bằng Microsoft Word.")

    doc.add_heading("DANH MỤC HÌNH ẢNH", level=1)
    figures = [
        "Hình 1. Sơ đồ Use Case tổng quát",
        "Hình 2. Kiến trúc ứng dụng Study Planner",
        "Hình 3. Sơ đồ quan hệ dữ liệu ERD",
        "Hình 4. Luồng hoạt động chính của ứng dụng",
        "Hình 5. Mô phỏng bố cục các màn hình chính",
    ]
    add_numbered(doc, figures)

    doc.add_heading("DANH MỤC BẢNG BIỂU", level=1)
    tables = [
        "Bảng 1. Thành phần và công nghệ sử dụng",
        "Bảng 2. Yêu cầu chức năng",
        "Bảng 3. Yêu cầu phi chức năng",
        "Bảng 4. Thiết kế bảng users",
        "Bảng 5. Thiết kế bảng courses",
        "Bảng 6. Thiết kế bảng tasks",
        "Bảng 7. Ánh xạ chức năng và lớp hiện thực",
        "Bảng 8. Bộ kiểm thử chức năng",
        "Bảng 9. Kết quả kiểm tra kỹ thuật",
    ]
    add_numbered(doc, tables)

    doc.add_heading("TÓM TẮT ĐỒ ÁN", level=1)
    add_body(doc, "Đề tài xây dựng ứng dụng Study Planner nhằm hỗ trợ sinh viên quản lý kế hoạch học tập cá nhân trên thiết bị Android. Bài toán tập trung vào việc tổ chức môn học, ghi nhận các nội dung cần học, xác định thời gian và mức ưu tiên, theo dõi tiến độ, xem lịch và nhận thông báo nhắc học.")
    add_body(doc, "Ứng dụng được phát triển bằng Java 17 trong Android Studio, giao diện sử dụng XML Layout và Material Components. Dữ liệu được lưu cục bộ bằng SQLite thông qua SQLiteOpenHelper; phiên đăng nhập được lưu bằng SharedPreferences. RecyclerView được dùng để hiển thị danh sách môn học và kế hoạch. AlarmManager, BroadcastReceiver và Notification được kết hợp để lập lịch nhắc học; BootReceiver khôi phục các lịch nhắc sau khi thiết bị khởi động lại.")
    add_body(doc, "Hệ thống gồm ba thực thể chính: người dùng, môn học và kế hoạch học tập. Mỗi kế hoạch có tiêu đề, ghi chú, ngày giờ, môn học, mức ưu tiên, thời lượng dự kiến, trạng thái và tùy chọn nhắc lịch. Ứng dụng hỗ trợ đăng ký, đăng nhập, CRUD môn học, CRUD kế hoạch, tìm kiếm, lọc, lịch học, thống kê và dark mode. Kết quả kiểm tra cho thấy project build thành công bằng Gradle và Android Lint không có lỗi.")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run("Từ khóa: ").bold = True
    p.add_run("Android, Java, XML Layout, SQLite, quản lý học tập, AlarmManager, Notification.")

    doc.add_heading("MỞ ĐẦU", level=1)
    doc.add_heading("1. Lý do chọn đề tài", level=2)
    add_body(doc, "Sinh viên thường phải đồng thời theo dõi nhiều môn học, bài tập, buổi ôn tập, đồ án và thời hạn nộp bài. Khi thông tin được ghi rải rác trong sổ tay, tin nhắn hoặc nhiều ứng dụng khác nhau, người học dễ quên nhiệm vụ, phân bổ thời gian chưa hợp lý và khó đánh giá tiến độ.")
    add_body(doc, "Điện thoại Android là thiết bị được sinh viên sử dụng hằng ngày, vì vậy một ứng dụng quản lý học tập hoạt động trực tiếp trên điện thoại có tính thực tiễn cao. Giải pháp cục bộ bằng SQLite không phụ thuộc kết nối Internet, phù hợp phạm vi môn học và vẫn thể hiện đầy đủ các nội dung kỹ thuật như tài khoản, quan hệ dữ liệu, CRUD, tìm kiếm, lịch, thông báo và thống kê.")
    doc.add_heading("2. Mục đích nghiên cứu", level=2)
    add_bullets(doc, [
        "Xây dựng ứng dụng Android giúp sinh viên tổ chức môn học và kế hoạch học tập tại một nơi thống nhất.",
        "Áp dụng Java, XML Layout, SQLite và các thành phần Android vào một bài toán có quan hệ dữ liệu rõ ràng.",
        "Hỗ trợ sinh viên xác định ưu tiên, ước lượng thời lượng và theo dõi mức độ hoàn thành.",
        "Tạo lịch nhắc học bằng AlarmManager và Notification để hạn chế quên kế hoạch.",
        "Rèn luyện quy trình khảo sát, phân tích, thiết kế, cài đặt, kiểm thử và viết tài liệu kỹ thuật.",
    ])
    doc.add_heading("3. Đối tượng nghiên cứu", level=2)
    add_body(doc, "Đối tượng sử dụng trực tiếp là sinh viên có nhu cầu quản lý môn học và lịch tự học cá nhân. Đối tượng kỹ thuật gồm nền tảng Android, ngôn ngữ Java, XML Layout, cơ sở dữ liệu SQLite, SharedPreferences, RecyclerView, Material Components và cơ chế thông báo của Android.")
    doc.add_heading("4. Phạm vi nghiên cứu", level=2)
    doc.add_heading("4.1. Phạm vi người dùng", level=3)
    add_body(doc, "Ứng dụng phục vụ một người dùng trên một thiết bị. Mỗi tài khoản chỉ truy cập dữ liệu môn học và kế hoạch gắn với user_id của mình.")
    doc.add_heading("4.2. Phạm vi chức năng", level=3)
    add_bullets(doc, [
        "Đăng ký, đăng nhập, lưu phiên và đăng xuất.",
        "Thêm, sửa, xóa và hiển thị môn học.",
        "Thêm, sửa, xóa, tìm kiếm, lọc và đánh dấu hoàn thành kế hoạch học tập.",
        "Chọn môn học, ngày giờ, mức ưu tiên, thời lượng và trạng thái nhắc lịch.",
        "Xem kế hoạch theo ngày trên lịch.",
        "Thống kê tổng kế hoạch, kế hoạch hoàn thành, kế hoạch còn lại, số môn và tổng giờ dự kiến.",
        "Hiển thị thông báo nhắc học và khôi phục lịch nhắc sau khi khởi động lại thiết bị.",
    ])
    add_body(doc, "Các chức năng đồng bộ đám mây, làm việc nhóm, chia sẻ lịch học, đăng nhập mạng xã hội và đồng bộ đa thiết bị chưa thuộc phạm vi phiên bản hiện tại.")

    # Chapter 1
    doc.add_heading("CHƯƠNG 1: TỔNG QUAN", level=1)
    doc.add_heading("1.1. Bối cảnh bài toán", level=2)
    add_body(doc, "Quản lý thời gian học tập không chỉ là ghi lại việc cần làm mà còn liên quan đến việc phân loại theo môn, xác định mức ưu tiên, ước lượng thời lượng và theo dõi tiến độ. Một danh sách công việc chung không phản ánh đầy đủ mối quan hệ giữa kế hoạch và môn học, vì vậy đề tài lựa chọn mô hình dữ liệu chuyên biệt cho sinh viên.")
    doc.add_heading("1.2. Vấn đề cần giải quyết", level=2)
    add_bullets(doc, [
        "Kế hoạch học tập phân tán, khó tìm kiếm và dễ bỏ sót.",
        "Sinh viên khó biết tổng khối lượng học tập và phần việc còn tồn đọng.",
        "Thiếu cơ chế nhắc đúng thời điểm và theo từng môn học.",
        "Các ứng dụng lớn có nhiều chức năng không cần thiết hoặc yêu cầu tài khoản trực tuyến.",
    ])
    doc.add_heading("1.3. Giải pháp đề xuất", level=2)
    add_body(doc, "Study Planner cung cấp luồng sử dụng ngắn: đăng nhập, tạo môn học, thêm kế hoạch và theo dõi trên trang tổng quan hoặc lịch. Giao diện sử dụng một màu chủ đạo xanh ngọc, các thành phần Material thống nhất, trạng thái rỗng rõ ràng và xác nhận trước thao tác xóa.")
    doc.add_heading("1.4. Kết quả và đóng góp", level=2)
    add_body(doc, "Sản phẩm hiện thực được một ứng dụng Android hoạt động offline với cơ sở dữ liệu quan hệ, xác thực cục bộ, nhiều màn hình, CRUD, lọc, thống kê và notification. Cấu trúc source được tách theo data/model, data/local, activity, fragment, adapter, notification và utils, giúp dễ trình bày và bảo trì.")

    # Chapter 2
    doc.add_heading("CHƯƠNG 2: NGHIÊN CỨU LÝ THUYẾT", level=1)
    doc.add_heading("2.1. Nền tảng Android và Java", level=2)
    add_body(doc, "Android là nền tảng di động dựa trên Linux, cung cấp bộ Android SDK để xây dựng Activity, Fragment, Service, BroadcastReceiver và các thành phần giao diện. Project sử dụng Java 17, compileSdk 35, targetSdk 35 và minSdk 23. Java được lựa chọn vì phù hợp nội dung học phần, có kiểu dữ liệu rõ ràng và được Android SDK hỗ trợ ổn định.")
    doc.add_heading("2.2. XML Layout và Material Components", level=2)
    add_body(doc, "Giao diện được định nghĩa bằng XML để tách trình bày khỏi xử lý nghiệp vụ. Material Components cung cấp TextInputLayout, MaterialButton, MaterialCardView, ChipGroup, BottomNavigationView và LinearProgressIndicator. Hệ thống style dùng chung giúp đồng nhất màu sắc, khoảng cách, trạng thái lỗi và dark mode.")
    doc.add_heading("2.3. Activity, Fragment và Intent", level=2)
    add_body(doc, "Activity đại diện cho các luồng độc lập như đăng nhập, đăng ký, tạo kế hoạch và quản lý môn học. MainActivity chứa các Fragment Tổng quan, Kế hoạch, Lịch và Cá nhân. Intent được dùng để mở màn hình và truyền dữ liệu kế hoạch cần chỉnh sửa.")
    doc.add_heading("2.4. RecyclerView và Adapter", level=2)
    add_body(doc, "RecyclerView hiển thị danh sách kế hoạch và môn học với cơ chế tái sử dụng ViewHolder. TaskAdapter ánh xạ StudyPlan sang item_task.xml; CourseAdapter ánh xạ Course sang item_course.xml. Adapter phát sự kiện click và thay đổi trạng thái về Fragment hoặc Activity xử lý.")
    doc.add_heading("2.5. SQLite và SQLiteOpenHelper", level=2)
    add_body(doc, "SQLite là hệ quản trị cơ sở dữ liệu quan hệ nhúng trong Android. DatabaseHelper chịu trách nhiệm tạo bảng, tạo chỉ mục, nâng phiên bản schema và cung cấp các phương thức truy vấn. Phiên bản cơ sở dữ liệu hiện tại là 3; migration bổ sung bảng courses và các cột priority, duration_minutes, reminder_enabled mà không xóa dữ liệu cũ.")
    doc.add_heading("2.6. SharedPreferences và bảo vệ mật khẩu", level=2)
    add_body(doc, "SessionManager lưu trạng thái đăng nhập, user_id và username bằng SharedPreferences. Mật khẩu không lưu trực tiếp mà được băm SHA-256 trước khi ghi vào SQLite. Cách này phù hợp phạm vi học tập, tuy nhiên ứng dụng thực tế nên sử dụng thuật toán có salt như PBKDF2, bcrypt hoặc Argon2.")
    doc.add_heading("2.7. AlarmManager, BroadcastReceiver và Notification", level=2)
    add_body(doc, "ReminderScheduler dùng AlarmManager.setAndAllowWhileIdle để lập lịch theo ngày giờ của kế hoạch. ReminderReceiver tạo notification channel và hiển thị thông báo. Trên Android 13 trở lên, ứng dụng yêu cầu quyền POST_NOTIFICATIONS. BootReceiver nhận BOOT_COMPLETED và lập lại các lịch nhắc còn hiệu lực sau khi thiết bị khởi động lại.")
    doc.add_heading("2.8. Thành phần và công nghệ sử dụng", level=2)
    add_table(doc,
              ["Thành phần", "Vai trò", "Ứng dụng trong đề tài"],
              [
                  ["Java 17", "Ngôn ngữ lập trình", "Xử lý nghiệp vụ, dữ liệu và lifecycle"],
                  ["XML Layout", "Định nghĩa giao diện", "Tất cả màn hình và item danh sách"],
                  ["SQLiteOpenHelper", "Quản lý cơ sở dữ liệu", "Users, courses và study plans"],
                  ["RecyclerView", "Danh sách hiệu năng cao", "Kế hoạch và môn học"],
                  ["Material Components", "Thành phần UI hiện đại", "Input, button, card, chip, navigation"],
                  ["SharedPreferences", "Lưu dữ liệu key-value", "Phiên đăng nhập"],
                  ["AlarmManager", "Lập lịch hệ thống", "Nhắc học theo ngày giờ"],
                  ["BroadcastReceiver", "Nhận sự kiện hệ thống", "Thông báo và khôi phục sau reboot"],
                  ["ExecutorService", "Thực thi nền", "Truy vấn SQLite không khóa UI"],
              ])
    add_caption(doc, "Bảng 1. Thành phần và công nghệ sử dụng")

    # Chapter 3
    doc.add_heading("CHƯƠNG 3: HIỆN THỰC HÓA NGHIÊN CỨU", level=1)
    doc.add_heading("3.1. Khảo sát và đặc tả yêu cầu", level=2)
    add_body(doc, "Từ nhu cầu thực tế của sinh viên, hệ thống được đặc tả theo các nhóm chức năng tài khoản, môn học, kế hoạch, lịch, nhắc học và thống kê. Các chức năng phải hoạt động offline, dữ liệu của các tài khoản phải tách biệt và thao tác thường dùng không vượt quá vài bước.")
    functional_rows = [
        ["F01", "Đăng ký tài khoản", "Kiểm tra username, email và mật khẩu trước khi lưu"],
        ["F02", "Đăng nhập và lưu phiên", "Mở màn hình chính khi thông tin hợp lệ"],
        ["F03", "CRUD môn học", "Tên môn bắt buộc; mã môn, giảng viên và màu tùy chọn"],
        ["F04", "CRUD kế hoạch học tập", "Lưu nội dung, môn, ưu tiên, thời lượng, ngày giờ"],
        ["F05", "Tìm kiếm và lọc", "Theo từ khóa, trạng thái và môn học"],
        ["F06", "Đánh dấu hoàn thành", "Cập nhật trạng thái và hủy reminder nếu cần"],
        ["F07", "Lịch học", "Hiển thị kế hoạch theo ngày được chọn"],
        ["F08", "Nhắc lịch", "Thông báo khi đến thời gian học"],
        ["F09", "Thống kê", "Đếm kế hoạch, môn học và tổng giờ dự kiến"],
        ["F10", "Hồ sơ và đăng xuất", "Hiển thị thông tin tài khoản và kết thúc phiên"],
    ]
    add_table(doc, ["Mã", "Yêu cầu", "Mô tả"], functional_rows, [1.5, 4.5, 10])
    add_caption(doc, "Bảng 2. Yêu cầu chức năng")
    add_table(doc, ["Nhóm", "Yêu cầu"],
              [
                  ["Hiệu năng", "Truy vấn SQLite thực hiện ngoài main thread; danh sách dùng RecyclerView."],
                  ["Ổn định", "Kiểm tra null/rỗng, lifecycle Fragment và session trước thao tác dữ liệu."],
                  ["Bảo mật", "Mật khẩu băm; dữ liệu truy vấn luôn ràng buộc theo user_id."],
                  ["Khả dụng", "Có validation inline, trạng thái rỗng, loading và xác nhận xóa."],
                  ["Tương thích", "Hoạt động từ Android API 23; hỗ trợ notification permission API 33+."],
                  ["Bảo trì", "Source tách package theo trách nhiệm và sử dụng resource dùng chung."],
              ])
    add_caption(doc, "Bảng 3. Yêu cầu phi chức năng")

    doc.add_heading("3.2. Phân tích Use Case", level=2)
    add_figure(doc, use_case, "Hình 1. Sơ đồ Use Case tổng quát", 6.6)
    add_body(doc, "Tác nhân chính là sinh viên. Sau khi đăng nhập, sinh viên có thể quản lý môn học và kế hoạch, tìm kiếm/lọc, xem lịch, đánh dấu hoàn thành, nhận nhắc lịch, xem thống kê và quản lý hồ sơ. Chức năng thêm hoặc cập nhật kế hoạch bao gồm lựa chọn môn học, ưu tiên, thời lượng và thời gian.")

    doc.add_heading("3.3. Thiết kế kiến trúc và cấu trúc source", level=2)
    add_figure(doc, architecture, "Hình 2. Kiến trúc ứng dụng Study Planner", 6.6)
    add_body(doc, "Ứng dụng sử dụng kiến trúc phân lớp đơn giản, phù hợp quy mô đồ án. UI chỉ chịu trách nhiệm thu thập dữ liệu và hiển thị trạng thái; DatabaseHelper tập trung logic SQLite; model biểu diễn dữ liệu; notification package xử lý lịch nhắc. ExecutorService ngăn truy vấn cơ sở dữ liệu làm giật giao diện.")
    tree = """com.example.personalplanner
├── activity/       Login, Register, Main, StudyPlan, Course
├── adapter/        TaskAdapter, CourseAdapter
├── data/
│   ├── local/      DatabaseHelper
│   └── model/      User, Course, StudyPlan, StudyStatistics
├── fragment/       Home, Task, Calendar, Profile
├── notification/   ReminderScheduler, ReminderReceiver, BootReceiver
└── utils/          SessionManager, PasswordUtils"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(tree)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9.5)

    doc.add_heading("3.4. Thiết kế cơ sở dữ liệu", level=2)
    add_figure(doc, erd, "Hình 3. Sơ đồ quan hệ dữ liệu ERD", 6.7)
    add_body(doc, "Một người dùng có nhiều môn học và nhiều kế hoạch. Một môn học có thể được gắn với nhiều kế hoạch. Trường course_id của kế hoạch cho phép giá trị 0 để biểu diễn trạng thái “Chưa phân loại”; khi xóa môn học, các kế hoạch liên quan được chuyển về trạng thái này thay vì bị xóa.")
    add_table(doc, ["Trường", "Kiểu", "Ràng buộc", "Ý nghĩa"],
              [
                  ["user_id", "INTEGER", "PK, AUTOINCREMENT", "Định danh người dùng"],
                  ["username", "TEXT", "NOT NULL, UNIQUE", "Tên đăng nhập"],
                  ["email", "TEXT", "NOT NULL, UNIQUE", "Email"],
                  ["password", "TEXT", "NOT NULL", "Mật khẩu đã băm SHA-256"],
                  ["created_at", "TEXT", "NOT NULL", "Thời điểm tạo tài khoản"],
              ])
    add_caption(doc, "Bảng 4. Thiết kế bảng users")
    add_table(doc, ["Trường", "Kiểu", "Ràng buộc", "Ý nghĩa"],
              [
                  ["course_id", "INTEGER", "PK, AUTOINCREMENT", "Định danh môn học"],
                  ["course_name", "TEXT", "NOT NULL", "Tên môn học"],
                  ["course_code", "TEXT", "Tùy chọn", "Mã môn học"],
                  ["lecturer", "TEXT", "Tùy chọn", "Tên giảng viên"],
                  ["color", "TEXT", "DEFAULT", "Màu nhận diện"],
                  ["user_id", "INTEGER", "FK, NOT NULL", "Chủ sở hữu môn học"],
              ])
    add_caption(doc, "Bảng 5. Thiết kế bảng courses")
    add_table(doc, ["Trường", "Kiểu", "Ràng buộc", "Ý nghĩa"],
              [
                  ["task_id", "INTEGER", "PK, AUTOINCREMENT", "Định danh kế hoạch"],
                  ["title", "TEXT", "NOT NULL", "Nội dung cần học"],
                  ["description", "TEXT", "Tùy chọn", "Ghi chú"],
                  ["date / time", "TEXT", "NOT NULL", "Ngày và giờ thực hiện"],
                  ["status", "INTEGER", "0 hoặc 1", "Chưa xong / hoàn thành"],
                  ["course_id", "INTEGER", "DEFAULT 0", "Môn học liên quan"],
                  ["priority", "INTEGER", "0..2", "Thấp, trung bình, cao"],
                  ["duration_minutes", "INTEGER", "DEFAULT 60", "Thời lượng dự kiến"],
                  ["reminder_enabled", "INTEGER", "0 hoặc 1", "Bật nhắc lịch"],
                  ["user_id", "INTEGER", "FK, NOT NULL", "Chủ sở hữu kế hoạch"],
              ])
    add_caption(doc, "Bảng 6. Thiết kế bảng tasks")

    doc.add_heading("3.5. Luồng điều hướng và xử lý", level=2)
    add_figure(doc, flow, "Hình 4. Luồng hoạt động chính của ứng dụng", 6.6)
    add_body(doc, "LoginActivity là launcher activity. Nếu SessionManager xác nhận đã đăng nhập, ứng dụng chuyển trực tiếp đến MainActivity. BottomNavigationView điều hướng giữa bốn Fragment; nút Thêm mở AddTaskActivity như một hành động, không phải tab cố định.")
    doc.add_heading("3.6. Hiện thực các module", level=2)
    mapping_rows = [
        ["Tài khoản", "LoginActivity, RegisterActivity, SessionManager, PasswordUtils"],
        ["Môn học", "CourseListActivity, CourseFormActivity, CourseAdapter"],
        ["Kế hoạch", "AddTaskActivity, TaskDetailActivity, TaskFragment, TaskAdapter"],
        ["Lịch", "CalendarFragment, CalendarView"],
        ["Thống kê", "HomeFragment, StudyStatistics"],
        ["Thông báo", "ReminderScheduler, ReminderReceiver, BootReceiver"],
        ["Dữ liệu", "DatabaseHelper, User, Course, StudyPlan"],
    ]
    add_table(doc, ["Module", "Lớp hiện thực"], mapping_rows)
    add_caption(doc, "Bảng 7. Ánh xạ chức năng và lớp hiện thực")
    doc.add_heading("3.6.1. Xác thực và phiên đăng nhập", level=3)
    add_body(doc, "Đăng ký kiểm tra độ dài username, định dạng email, độ dài mật khẩu và xác nhận mật khẩu. DatabaseHelper kiểm tra trùng username/email không phân biệt hoa thường. Khi đăng nhập thành công, user_id và username được lưu vào SharedPreferences và dùng để giới hạn mọi truy vấn dữ liệu.")
    doc.add_heading("3.6.2. Quản lý môn học", level=3)
    add_body(doc, "Sinh viên nhập tên môn, mã môn, giảng viên và chọn màu nhận diện. Danh sách môn học được sắp xếp theo tên. Khi xóa môn, DatabaseHelper thực hiện transaction: cập nhật course_id của kế hoạch về 0 rồi xóa bản ghi môn học.")
    doc.add_heading("3.6.3. Quản lý kế hoạch học tập", level=3)
    add_body(doc, "Form kế hoạch yêu cầu nội dung và thời lượng lớn hơn 0. Người dùng chọn môn học, mức ưu tiên, ngày giờ và bật nhắc lịch. Danh sách hỗ trợ tìm trong title, description và course_name; bộ lọc gồm trạng thái và môn học.")
    doc.add_heading("3.6.4. Lịch và thống kê", level=3)
    add_body(doc, "CalendarFragment truy vấn kế hoạch theo ngày. HomeFragment sử dụng câu truy vấn tổng hợp để tính tổng kế hoạch, số hoàn thành, số còn lại, số môn và tổng phút dự kiến; phần trăm hoàn thành được hiển thị bằng LinearProgressIndicator.")
    doc.add_heading("3.6.5. Nhắc lịch học", level=3)
    add_body(doc, "Khi lưu kế hoạch có bật nhắc, ứng dụng tạo PendingIntent có requestCode bằng planId và đăng ký AlarmManager. Khi kế hoạch hoàn thành hoặc bị xóa, lịch nhắc được hủy. ReminderReceiver hiển thị tên kế hoạch và môn học; BootReceiver khôi phục reminder sau khi khởi động lại máy.")
    doc.add_heading("3.7. Thiết kế giao diện", level=2)
    add_body(doc, "Bảng màu chủ đạo gồm xanh ngọc #1F6F68, cam đất #D08A4B, nền #F6F7F4 và chữ #17201E. Card bo góc 18dp, nút cao tối thiểu 52dp, TextInputLayout có lỗi inline và khoảng cách chính 14–24dp. values-night cung cấp bảng màu tối tương ứng.")

    # Chapter 4
    doc.add_heading("CHƯƠNG 4: KẾT QUẢ NGHIÊN CỨU", level=1)
    doc.add_heading("4.1. Môi trường xây dựng và kiểm tra", level=2)
    add_table(doc, ["Hạng mục", "Thông số"],
              [
                  ["IDE", "Android Studio"],
                  ["Ngôn ngữ", "Java 17"],
                  ["Giao diện", "XML Layout, Material Components 1.12.0"],
                  ["Cơ sở dữ liệu", "SQLiteOpenHelper, schema version 3"],
                  ["Android SDK", "compileSdk 35, targetSdk 35, minSdk 23"],
                  ["Build", "Gradle assembleDebug"],
                  ["Phân tích tĩnh", "Android Lint"],
              ])
    doc.add_heading("4.2. Kết quả giao diện", level=2)
    add_figure(doc, ui_mockups, "Hình 5. Mô phỏng bố cục các màn hình chính theo XML hiện thực", 6.8)
    add_body(doc, "Màn hình đăng nhập tập trung vào hai trường thông tin và liên kết đăng ký. Trang tổng quan hiển thị tiến độ, số lượng kế hoạch, số môn và tổng giờ dự kiến. Danh sách kế hoạch thể hiện môn học, nội dung, ngày giờ, ưu tiên, thời lượng và checkbox trạng thái. Các form sử dụng cùng hệ thống style để bảo đảm tính nhất quán.")
    doc.add_heading("4.3. Kiểm thử chức năng", level=2)
    test_rows = [
        ["TC01", "Đăng ký hợp lệ", "Nhập username, email và mật khẩu hợp lệ", "Tài khoản được tạo"],
        ["TC02", "Email sai", "Nhập email không đúng định dạng", "Hiện lỗi tại ô email"],
        ["TC03", "Tài khoản trùng", "Dùng username hoặc email đã tồn tại", "Từ chối và thông báo lỗi"],
        ["TC04", "Đăng nhập đúng", "Nhập đúng username/mật khẩu", "Mở MainActivity"],
        ["TC05", "Đăng nhập sai", "Nhập sai mật khẩu", "Không mở ứng dụng chính"],
        ["TC06", "Thêm môn học", "Nhập tên, mã và giảng viên", "Môn xuất hiện trong danh sách"],
        ["TC07", "Sửa môn học", "Chỉnh nội dung và lưu", "Dữ liệu được cập nhật"],
        ["TC08", "Xóa môn học", "Xác nhận xóa", "Môn bị xóa, kế hoạch không mất"],
        ["TC09", "Thêm kế hoạch", "Chọn môn, ưu tiên, thời lượng và ngày giờ", "Kế hoạch được lưu"],
        ["TC10", "Thời lượng sai", "Nhập 0 hoặc ký tự", "Hiện lỗi, không lưu"],
        ["TC11", "Tìm kiếm", "Nhập từ khóa title/mô tả/môn", "Chỉ hiện kết quả phù hợp"],
        ["TC12", "Lọc trạng thái", "Chọn Chưa xong/Hoàn thành", "Danh sách lọc đúng"],
        ["TC13", "Lọc môn học", "Chọn một môn trong Spinner", "Chỉ hiện kế hoạch của môn"],
        ["TC14", "Đánh dấu hoàn thành", "Chọn checkbox", "Trạng thái cập nhật, reminder hủy"],
        ["TC15", "Xem lịch", "Chọn một ngày", "Hiện kế hoạch của ngày đó"],
        ["TC16", "Thông báo", "Bật reminder với thời gian tương lai", "Notification xuất hiện đúng lịch"],
        ["TC17", "Khởi động lại", "Reboot thiết bị trước giờ nhắc", "BootReceiver lập lại reminder"],
        ["TC18", "Tách dữ liệu", "Đăng nhập hai tài khoản", "Mỗi tài khoản chỉ thấy dữ liệu của mình"],
    ]
    add_table(doc, ["Mã", "Chức năng", "Thao tác", "Kết quả mong muốn"], test_rows,
              [1.2, 3.5, 6.2, 6.2])
    add_caption(doc, "Bảng 8. Bộ kiểm thử chức năng")
    doc.add_heading("4.4. Kết quả kiểm tra kỹ thuật", level=2)
    add_table(doc, ["Kiểm tra", "Kết quả", "Đánh giá"],
              [
                  ["clean assembleDebug", "BUILD SUCCESSFUL", "APK debug được tạo thành công"],
                  ["Android Lint", "0 errors, 4 warnings", "Không có lỗi; warning thuộc phiên bản SDK/dependency"],
                  ["Resource linking", "Thành công", "XML, style, string và manifest hợp lệ"],
                  ["Java compilation", "Thành công", "Package và API mới biên dịch đúng"],
                  ["APK đầu ra", "app-debug.apk", "Có thể cài trên thiết bị API 23+"],
              ])
    add_caption(doc, "Bảng 9. Kết quả kiểm tra kỹ thuật")
    doc.add_heading("4.5. Đánh giá kết quả", level=2)
    add_body(doc, "Ứng dụng đáp ứng các yêu cầu cốt lõi của đề tài: tài khoản, cơ sở dữ liệu quan hệ, CRUD nhiều thực thể, tìm kiếm/lọc, lịch, thông báo và thống kê tiến độ. Giao diện có dark mode, validation, empty state và loading. Cấu trúc source đã được tối ưu theo trách nhiệm thay vì dồn logic vào một Activity.")
    add_body(doc, "Hạn chế hiện tại là dữ liệu chỉ lưu trên một thiết bị; thông báo dùng AlarmManager không phải exact alarm nên có thể trễ nhẹ do cơ chế tiết kiệm pin; mật khẩu sử dụng SHA-256 chưa có salt; chưa có unit test và instrumentation test tự động.")

    # Chapter 5
    doc.add_heading("CHƯƠNG 5: KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", level=1)
    doc.add_heading("5.1. Kết luận", level=2)
    add_body(doc, "Đồ án đã xây dựng thành công ứng dụng Study Planner dành cho sinh viên trên nền tảng Android. Sản phẩm giúp tổ chức môn học, lập kế hoạch, theo dõi tiến độ và nhận nhắc lịch. Qua quá trình thực hiện, nhóm đã vận dụng được kiến thức Java, XML, SQLite, Activity, Fragment, RecyclerView, SharedPreferences, AlarmManager và BroadcastReceiver.")
    add_body(doc, "Việc chuyển từ mô hình to-do tổng quát sang mô hình học tập giúp đề tài có quan hệ dữ liệu và nghiệp vụ rõ ràng hơn. Ba thực thể User, Course và StudyPlan đủ để thể hiện quan hệ một-nhiều, CRUD và truy vấn tổng hợp nhưng vẫn phù hợp thời gian triển khai môn học.")
    doc.add_heading("5.2. Hướng phát triển", level=2)
    add_bullets(doc, [
        "Chuyển lớp dữ liệu sang Room Database và Repository để tăng khả năng kiểm thử.",
        "Đồng bộ tài khoản và dữ liệu bằng Firebase hoặc REST API.",
        "Bổ sung lịch học lặp theo tuần, deadline, học kỳ và tín chỉ.",
        "Thêm biểu đồ tiến độ theo môn, tuần và tháng.",
        "Hỗ trợ import/export lịch bằng định dạng iCalendar.",
        "Sử dụng WorkManager cho các tác vụ nền dài hạn và đồng bộ.",
        "Nâng cấp bảo mật mật khẩu bằng PBKDF2/Argon2 và Android Keystore.",
        "Bổ sung unit test, database migration test và UI test bằng Espresso.",
        "Tối ưu giao diện cho tablet và hỗ trợ đa ngôn ngữ.",
    ])

    doc.add_heading("DANH MỤC TÀI LIỆU THAM KHẢO", level=1)
    references = [
        "Android Developers. (2026). Application fundamentals. https://developer.android.com/guide/components/fundamentals",
        "Android Developers. (2026). Save data using SQLite. https://developer.android.com/training/data-storage/sqlite",
        "Android Developers. (2026). Create a notification. https://developer.android.com/develop/ui/views/notifications/build-notification",
        "Android Developers. (2026). Schedule alarms. https://developer.android.com/develop/background-work/services/alarms",
        "Android Developers. (2026). RecyclerView. https://developer.android.com/develop/ui/views/layout/recyclerview",
        "Google. Material Design 3. https://m3.material.io/",
        "Phillips, B., Stewart, C., Marsicano, K. (2022). Android Programming: The Big Nerd Ranch Guide. Big Nerd Ranch.",
        "SQLite Documentation. https://www.sqlite.org/docs.html",
    ]
    add_numbered(doc, references)

    doc.add_heading("PHỤ LỤC", level=1)
    doc.add_heading("A. Cách build project", level=2)
    add_body(doc, "Mở project trong Android Studio, chọn Gradle JDK 17 và thực hiện Sync Project with Gradle Files. Có thể build bằng PowerShell:")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(".\\gradlew.bat clean assembleDebug\n.\\gradlew.bat lintDebug")
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    doc.add_heading("B. File APK", level=2)
    add_body(doc, "APK debug sau khi build được tạo tại app/build/outputs/apk/debug/app-debug.apk.")
    doc.add_heading("C. Quyền Android", level=2)
    add_bullets(doc, [
        "POST_NOTIFICATIONS: cho phép hiển thị notification trên Android 13 trở lên.",
        "RECEIVE_BOOT_COMPLETED: khôi phục lịch nhắc học sau khi thiết bị khởi động.",
    ])
    doc.add_heading("D. Giả định triển khai", level=2)
    add_body(doc, "Ứng dụng hoạt động offline; mỗi thiết bị lưu cơ sở dữ liệu riêng. Ngày và giờ được lưu theo định dạng yyyy-MM-dd và HH:mm để sắp xếp trực tiếp trong SQLite. Dữ liệu mẫu, ảnh mô phỏng và số liệu kiểm thử trong báo cáo phục vụ mục đích minh họa đồ án.")

    for section in doc.sections:
        configure_section(section)
        set_header_footer(section)
    doc.sections[0].different_first_page_header_footer = True
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    make_document()
